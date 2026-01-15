# ----------------------------
# ✅ IMPORTS
# ----------------------------
import streamlit as st
import pandas as pd
pd.options.mode.copy_on_write = True  # 🔑 evita cópias profundas desnecessárias

import plotly.express as px
from datetime import datetime
from fpdf import FPDF
import tempfile
import os
from PIL import Image
from sklearn.model_selection import train_test_split
from sklearn.ensemble import RandomForestClassifier
from sklearn.preprocessing import LabelEncoder
from sklearn.metrics import accuracy_score
import time
import plotly.io as pio
import kaleido  # força o loading do kaleido para evitar erro de exportação
import unicodedata
import numpy as np
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import requests
from io import BytesIO
from docx.shared import RGBColor
from docx.shared import Pt
from docx.shared import RGBColor, Pt
import gc

pio.kaleido.scope.default_format = "png"

# ----------------------------
# ✅ CONFIGURAÇÃO INICIAL DA PÁGINA
# ----------------------------
st.set_page_config(
    page_title="Dashboard Analítico",
    page_icon="https://raw.githubusercontent.com/carlinhosg7/streamlit02/main/logo_kidy_icon.ico",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ----------------------------
# 🔐 FUNÇÃO DE AUTENTICAÇÃO
# ----------------------------
def autenticar_usuario_excel(caminho_arquivo):
    try:
        df_usuarios = pd.read_excel(caminho_arquivo, engine="openpyxl")

        # Normaliza os nomes das colunas
        df_usuarios.columns = [
            unicodedata.normalize('NFKD', col).encode('ascii', errors='ignore').decode('utf-8').strip().lower().replace(" ", "_")
            for col in df_usuarios.columns
        ]

        # Converte os dados para string
        df_usuarios['usuario'] = df_usuarios['usuario'].apply(lambda x: str(x).strip())
        df_usuarios['senha'] = df_usuarios['senha'].astype(str).str.strip()

        # Inicializa autenticação
        st.session_state['autenticado'] = st.session_state.get('autenticado', False)

        if not st.session_state['autenticado']:
            with st.sidebar:
                st.markdown("### 🔐 Login")
                usuario = st.text_input("Usuário").strip()
                senha = st.text_input("Senha", type="password").strip()

                if st.button("Entrar"):
                    if usuario in df_usuarios['usuario'].values:
                        senha_valida = df_usuarios[df_usuarios['usuario'] == usuario]['senha'].values[0]
                        if senha == senha_valida:
                            st.session_state['autenticado'] = True
                            st.session_state['codigo_representante'] = usuario
                            st.success("✅ Login realizado com sucesso!")
                            st.rerun()
                        else:
                            st.error("❌ Senha incorreta.")
        if not st.session_state['autenticado']:
            st.stop()

    except Exception as e:
        st.error(f"Erro ao carregar planilha de autenticação: {e}")
        st.stop()

# ----------------------------
# 🚀 AUTENTICAÇÃO
# ----------------------------
autenticar_usuario_excel("auth.xlsx")

# ----------------------------
# 🖼️ LOGO
# ----------------------------
CAMINHO_LOGO_LOCAL = "C:/preditiva/streamlit02/logo_kidy.png"
logo_kidy = None
try:
    logo_kidy = Image.open(CAMINHO_LOGO_LOCAL)
except Exception as e:
    st.warning(f"⚠️ Não foi possível carregar a logo local: {e}")

st.title("📊 Dashboard Analítico")
if logo_kidy:
    st.sidebar.image(logo_kidy, width=100)
    st.image(logo_kidy, width=150)

# ----------------------------
# 🎨 CSS CUSTOMIZADO
# ----------------------------
def add_custom_css():
    st.markdown("""
        <style>
        body { background-color: #1e1e1e; color: #ffffff; }
        .block-container { padding: 2rem; }
        div.stButton > button:first-child {
            background-color: #E60012; color: white; border-radius: 8px;
            height: 3em; width: 100%; font-weight: bold; border: none; transition: 0.3s;
        }
        div.stButton > button:first-child:hover { background-color: #A3000B; }
        footer {visibility: hidden;}
        .metric-card { background-color: #2d2d2d; padding: 16px; border-radius: 10px;
            box-shadow: 1px 1px 8px rgba(0,0,0,0.3); text-align: center; margin-bottom: 8px;
            border: 1px solid #fba72033; }
        .metric-label { font-size: 13px; color: #bbbbbb; }
        .metric-value { font-size: 22px; font-weight: bold; color: #F7A400; }
        </style>
    """, unsafe_allow_html=True)

add_custom_css()

# ----------------------------
# 📆 DICIONÁRIO MESES
# ----------------------------
meses_portugues = {
    1: 'Janeiro', 2: 'Fevereiro', 3: 'Março', 4: 'Abril',
    5: 'Maio', 6: 'Junho', 7: 'Julho', 8: 'Agosto',
    9: 'Setembro', 10: 'Outubro', 11: 'Novembro', 12: 'Dezembro'
}

# ----------------------------
# 📁 Caminhos locais
# ----------------------------
CAMINHO_LOCAL = "C:/preditiva/streamlit02"
URLS_DADOS = [
    f"{CAMINHO_LOCAL}/DADOS_PREDITIVA_1.parquet",
    f"{CAMINHO_LOCAL}/DADOS_PREDITIVA_2.parquet",
    f"{CAMINHO_LOCAL}/DADOS_PREDITIVA_3.parquet",
    f"{CAMINHO_LOCAL}/DADOS_PREDITIVA_4.parquet",
    f"{CAMINHO_LOCAL}/DADOS_PREDITIVA_5.parquet",
    f"{CAMINHO_LOCAL}/DADOS_PREDITIVA_6.parquet"
]

COLUNAS_NECESSARIAS = [
    'Codigo Representante', 'Codigo Supervisor',
    'Codigo Grupo Cliente', 'Codigo Cliente',
    'Grupo Cliente', 'Razao Social', 'Linha',
    'Data Cadastro', 'Data Ultima Compra',
    'Qtd Venda', 'Vlr Venda'
]

# ----------------------------
# 🚀 CARREGAMENTO OTIMIZADO (4GB-friendly)
# ----------------------------
@st.cache_resource
def carregar_dados_processados():
    try:
        dfs_filtrados = []
        colunas_base = None

        cod_rep_login = str(st.session_state.get('codigo_representante', '')).strip()
        eh_admin = (cod_rep_login.lower() == 'admin' or cod_rep_login == '')

        login_digits = ''.join(ch for ch in cod_rep_login if ch.isdigit())
        cod_norm = login_digits.lstrip('0') or login_digits or cod_rep_login
        cand_str = {cod_rep_login, cod_norm}
        cand_int = int(cod_norm) if cod_norm.isdigit() else None
        cand_float = float(cand_int) if isinstance(cand_int, int) else None

        for url in URLS_DADOS:
            # lê só as colunas necessárias (pyarrow é mais leve)
            df_temp = pd.read_parquet(url, columns=COLUNAS_NECESSARIAS, engine="pyarrow")

            # filtro por representante sem converter a coluna inteira
            if not eh_admin and 'Codigo Representante' in df_temp.columns:
                s = df_temp['Codigo Representante']
                mask = pd.Series(False, index=s.index)

                if s.dtype.kind in ('O', 'U', 'S') or str(s.dtype).startswith(('string', 'category')):
                    mask |= s.isin(cand_str)
                    idx = (~mask) & s.notna()
                    if idx.any():
                        s_part = pd.Series(s[idx], dtype="string")
                        mask.loc[idx] = s_part.str.strip().str.lstrip('0').isin({cod_norm})

                if cand_int is not None and s.dtype.kind in ('i', 'u'):
                    mask |= (s == cand_int)

                if cand_float is not None and s.dtype.kind == 'f':
                    mask |= (s == cand_float)

                df_temp = df_temp.loc[mask]
                if df_temp.empty:
                    del df_temp
                    gc.collect()
                    continue

            # valida esquema
            if colunas_base is None:
                colunas_base = df_temp.columns
            elif not df_temp.columns.equals(colunas_base):
                st.warning(f"⚠️ Estrutura diferente detectada no arquivo: {url}")
                gc.collect()
                continue

            # tipagem leve
            for col in ['Codigo Supervisor', 'Codigo Grupo Cliente', 'Codigo Cliente',
                        'Grupo Cliente', 'Razao Social', 'Linha']:
                if col in df_temp.columns:
                    df_temp[col] = pd.Series(df_temp[col], dtype="string").str.strip()

            for col in ['Data Cadastro', 'Data Ultima Compra']:
                if col in df_temp.columns:
                    df_temp[col] = pd.to_datetime(df_temp[col], errors='coerce')

            if 'Qtd Venda' in df_temp.columns:
                df_temp['Qtd Venda'] = pd.to_numeric(df_temp['Qtd Venda'], errors='coerce').fillna(0).astype('int32')
            if 'Vlr Venda' in df_temp.columns:
                df_temp['Vlr Venda'] = pd.to_numeric(df_temp['Vlr Venda'], errors='coerce').fillna(0).astype('float32')

            dfs_filtrados.append(df_temp)
            del df_temp
            gc.collect()

        if not dfs_filtrados:
            return pd.DataFrame(columns=COLUNAS_NECESSARIAS)

        df = pd.concat(dfs_filtrados, ignore_index=True)
        del dfs_filtrados
        gc.collect()

        if 'Codigo Grupo Cliente' in df.columns:
            df['Codigo Grupo Cliente'] = pd.Series(df['Codigo Grupo Cliente'], dtype="string").str.upper()
        if 'Codigo Cliente' in df.columns:
            df['Codigo Cliente'] = pd.Series(df['Codigo Cliente'], dtype="string").str.upper()

        for c in ['Codigo Representante', 'Codigo Supervisor', 'Codigo Grupo Cliente',
                  'Codigo Cliente', 'Grupo Cliente', 'Razao Social', 'Linha']:
            if c in df.columns and df[c].nunique(dropna=True) <= 50000:
                df[c] = df[c].astype('category')

        # preço médio básico (ajustaremos novamente no filtro)
        if {'Vlr Venda', 'Qtd Venda'}.issubset(df.columns):
            qtd = df['Qtd Venda'].replace(0, pd.NA)
            df['Preço Médio Produto'] = (df['Vlr Venda'] / qtd).fillna(0).round(2).astype('float32')

        # mensagem de contexto
        cod_norm = ''.join(ch for ch in cod_rep_login if ch.isdigit()).lstrip('0') or cod_rep_login
        if eh_admin:
            st.markdown("🟢 **Modo Admin: Visualizando todos os dados**")
        else:
            st.markdown(f"🆔 **Representante:** `{(cod_norm or cod_rep_login).upper()}`")

        gc.collect()
        return df

    except Exception as e:
        import traceback
        st.error("❌ Erro ao carregar dados:")
        st.code(traceback.format_exc())
        return pd.DataFrame()

# ----------------------------
# 📥 CARREGA DADOS
# ----------------------------
df = carregar_dados_processados()
if df.empty:
    st.warning("⚠️ Sem dados para exibir com os filtros atuais.")
    st.stop()

# ----------------------------
# 🔎 OPÇÕES FORMATADAS
# ----------------------------
opcoes_grupo_cliente = df[['Codigo Grupo Cliente', 'Grupo Cliente']].drop_duplicates()
opcoes_grupo_cliente['Busca'] = (
    opcoes_grupo_cliente['Codigo Grupo Cliente'].astype(str) + ' - ' +
    opcoes_grupo_cliente['Grupo Cliente'].astype(str)
)
busca_grupo_lista = sorted(opcoes_grupo_cliente['Busca'].tolist())
busca_grupo = st.sidebar.selectbox("🔍 Buscar Grupo Cliente:", [''] + busca_grupo_lista)

opcoes_cliente = df[['Codigo Cliente', 'Razao Social']].drop_duplicates()
opcoes_cliente['Busca'] = (
    opcoes_cliente['Codigo Cliente'].astype(str) + ' - ' +
    opcoes_cliente['Razao Social'].astype(str)
)
busca_cliente_lista = sorted(opcoes_cliente['Busca'].tolist())
busca_cliente = st.sidebar.selectbox("🔍 Buscar Cliente:", [''] + busca_cliente_lista)

# ----------------------------
# 🧩 EXTRAÇÃO DE CÓDIGOS E PERÍODO
# ----------------------------
codigo_grupo_cliente = busca_grupo.split(' - ')[0].strip().upper() if busca_grupo else ''
codigo_cliente = busca_cliente.split(' - ')[0].strip().upper() if busca_cliente else ''

if 'Data Cadastro' in df.columns and df['Data Cadastro'].notna().any():
    data_min = df['Data Cadastro'].min().date()
    data_max = df['Data Cadastro'].max().date()
else:
    data_min = datetime(2000, 1, 1).date()
    data_max = datetime.today().date()

data_inicio_padrao = datetime(2024, 1, 1).date()
data_fim_padrao = min(datetime.today().date(), data_max)

periodo = st.sidebar.date_input(
    "Período da análise:",
    value=(data_inicio_padrao, data_fim_padrao),
    min_value=data_min,
    max_value=data_max
)

# ----------------------------
# 🔘 BOTÃO PRINCIPAL
# ----------------------------
if st.sidebar.button("🔎 Analisar Grupo/Cliente"):

    if not codigo_grupo_cliente and not codigo_cliente:
        st.sidebar.warning("⚠️ Informe pelo menos um código!")
    else:
        with st.spinner('🔎 Analisando dados...'):

            # ==== FILTROS INICIAIS (sem cópias profundas) ====
            filtro_inicial = df
            if codigo_cliente:
                filtro_inicial = filtro_inicial.loc[filtro_inicial['Codigo Cliente'] == codigo_cliente]
            elif codigo_grupo_cliente:
                filtro_inicial = filtro_inicial.loc[filtro_inicial['Codigo Grupo Cliente'] == codigo_grupo_cliente]

            dados_filtrados = filtro_inicial.loc[
                (filtro_inicial['Data Cadastro'] >= pd.to_datetime(periodo[0])) &
                (filtro_inicial['Data Cadastro'] <= pd.to_datetime(periodo[1]))
            ]
            if dados_filtrados.empty:
                st.warning("⚠️ Nenhum dado encontrado no período!")
                st.stop()

            # ==== TIPAGEM E CÁLCULOS VETORIZADOS ====
            dados_filtrados['Qtd Venda'] = pd.to_numeric(dados_filtrados['Qtd Venda'], errors='coerce').fillna(0).astype('int32')
            dados_filtrados['Vlr Venda'] = pd.to_numeric(dados_filtrados['Vlr Venda'], errors='coerce').fillna(0).astype('float32')

            mask_validos = (dados_filtrados['Qtd Venda'] >= 0) & (dados_filtrados['Vlr Venda'] >= 0)
            dados_filtrados = dados_filtrados.loc[mask_validos]

            qtd_np = dados_filtrados['Qtd Venda'].to_numpy()
            vlr_np = dados_filtrados['Vlr Venda'].to_numpy(dtype='float32')

            preco_medio_np = np.where(qtd_np > 0, vlr_np / qtd_np, 0).astype('float32')
            dados_filtrados['Preço Médio Produto'] = preco_medio_np

            mask_base = (qtd_np > 0) & (vlr_np > 0)
            if mask_base.any():
                media_preco_par = (vlr_np[mask_base] / qtd_np[mask_base]).mean()
            else:
                media_preco_par = 60.0
                st.warning(f"⚠️ Nenhuma venda com valor encontrada. Estimativa padrão de R$ {media_preco_par:.2f} por par foi aplicada.")

            dados_filtrados['Vlr Venda Corrigido'] = np.where(vlr_np > 0, vlr_np, qtd_np * media_preco_par).astype('float32')

            # ==== INFO DE CABEÇALHO ====
            dados_filtrados['Ano'] = dados_filtrados['Data Cadastro'].dt.year
            nome_grupo = dados_filtrados['Grupo Cliente'].iloc[0]
            total_lojas = dados_filtrados['Codigo Cliente'].nunique()

            mapa_supervisores = {
                '9902': 'Centro Oeste',
                '9907': 'Sul',
                '9914': 'Norte / Nordeste',
                '9915': 'REM',
                '9916': 'SPC',
                '9917': 'SPI'
            }

            cods_representantes = sorted(dados_filtrados['Codigo Representante'].astype(str).dropna().unique())
            cods_repr_str = ', '.join(cods_representantes)
            st.session_state["codigo_repr_para_arquivo"] = cods_representantes[0] if len(cods_representantes) else "rep"

            cod_supervisor = dados_filtrados['Codigo Supervisor'].astype(str).str.strip().dropna().unique()
            if len(cod_supervisor) > 0:
                supervisor_id = str(cod_supervisor[0]).split('.')[0]
                nome_supervisor = mapa_supervisores.get(supervisor_id, f"Código {supervisor_id} não mapeado")
            else:
                nome_supervisor = "Não informado"

            st.markdown(f"## 📍 Grupo Cliente: {nome_grupo} | 🏬 Lojas: {total_lojas} | 🆔 Representante(s): {cods_repr_str}")
            st.markdown(f"#### 🧑‍💼 Supervisor: {nome_supervisor}")

            ultima_data_compra = dados_filtrados['Data Ultima Compra'].max()
            ultima_compra = ultima_data_compra.strftime('%d/%m/%Y') if pd.notnull(ultima_data_compra) else 'Sem compras'
            primeira_data = dados_filtrados['Data Cadastro'].min()
            ultima_data = dados_filtrados['Data Cadastro'].max()
            periodo_analise = f"{primeira_data.strftime('%d/%m/%Y')} até {ultima_data.strftime('%d/%m/%Y')}"
            vendas_totais = int(dados_filtrados['Qtd Venda'].sum())
            melhor_mes_num = int(dados_filtrados['Data Cadastro'].dt.month.mode()[0])
            melhor_mes_nome = meses_portugues.get(melhor_mes_num, 'Mês inválido')

            col1, col2, col3 = st.columns(3)
            for col, label, value in zip(
                [col1, col2, col3],
                ['📅 Última Compra', '🕒 Período da Análise', '📈 Melhor Mês para Oferta'],
                [ultima_compra, periodo_analise, melhor_mes_nome]
            ):
                col.markdown(f"""
                    <div class="metric-card">
                        <div class="metric-label">{label}</div>
                        <div class="metric-value">{value}</div>
                    </div>
                """, unsafe_allow_html=True)
            st.success(f"📦 Total de Itens Vendidos: {vendas_totais:,} unidades")

            # ==== COLEÇÕES (vetorizado) ====
            def identificar_colecao_dt(series_dt):
                m = series_dt.dt.month
                a = series_dt.dt.year
                inverno = np.where(m >= 11, a + 1, a)
                rotulo = np.where((m >= 5) & (m <= 10), 'Verão ' + a.astype(str), 'Inverno ' + inverno.astype(str))
                return pd.Series(rotulo, index=series_dt.index)

            dados_filtrados = dados_filtrados.loc[dados_filtrados['Data Cadastro'].notna()]
            dados_filtrados['Colecao'] = identificar_colecao_dt(dados_filtrados['Data Cadastro'])

            grp = dados_filtrados.groupby('Colecao', as_index=False).agg(
                Qtd_Venda=('Qtd Venda', 'sum'),
                Vlr_Venda_Corrigido=('Vlr Venda Corrigido', 'sum'),
                Data_Inicial=('Data Cadastro', 'min'),
                Data_Final=('Data Cadastro', 'max')
            )
            grp['Ano'] = grp['Colecao'].str.extract(r'(\d{4})').astype(int)

            hoje = datetime.today()
            m_atual, a_atual = hoje.month, hoje.year
            colecao_vigente = f"Verão {a_atual}" if 5 <= m_atual <= 10 else (f"Inverno {a_atual+1}" if m_atual >= 11 else f"Inverno {a_atual}")
            if colecao_vigente not in set(grp['Colecao']):
                grp = pd.concat([
                    pd.DataFrame([{
                        'Colecao': colecao_vigente,
                        'Qtd_Venda': 0,
                        'Vlr_Venda_Corrigido': 0.0,
                        'Data_Inicial': hoje,
                        'Data_Final': hoje,
                        'Ano': int(colecao_vigente.split()[1])
                    }]),
                    grp
                ], ignore_index=True)

            colecoes_exibir = (
                grp.sort_values('Data_Final', ascending=False)
                   .drop_duplicates('Colecao')
                   .head(3)
                   .copy()
            )
            colecoes_exibir['Pares Vendidos'] = colecoes_exibir['Qtd_Venda'].astype(int)
            colecoes_exibir['Valor Vendido (R$)'] = colecoes_exibir['Vlr_Venda_Corrigido'].map(
                lambda x: f"R$ {x:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
            )
            colecoes_exibir['Período da Coleta'] = np.where(
                colecoes_exibir['Data_Inicial'].notna() & colecoes_exibir['Data_Final'].notna(),
                colecoes_exibir['Data_Inicial'].dt.strftime('%d/%m/%Y') + ' a ' + colecoes_exibir['Data_Final'].dt.strftime('%d/%m/%Y'),
                'Período inválido'
            )
            colecoes_exibir = colecoes_exibir[['Colecao','Pares Vendidos','Valor Vendido (R$)','Período da Coleta']]
            colecoes_exibir.columns = ['Coleção','Pares Vendidos','Valor Vendido (R$)','Período da Coleta']
            st.session_state["colecoes_exibir"] = colecoes_exibir

            st.markdown("### 👟 Vendas das 3 Últimas Coleções (Pares e Valores)")
            st.table(colecoes_exibir)

            # ==== ÚLTIMOS 12 MESES (vetorizado) ====
            dados_filtrados['AnoMes'] = dados_filtrados['Data Cadastro'].dt.to_period('M').dt.to_timestamp()
            ultimos_12 = (
                dados_filtrados.groupby('AnoMes', as_index=False)
                .agg(Qtd_Venda=('Qtd Venda','sum'), Vlr_Venda_Corrigido=('Vlr Venda Corrigido','sum'))
                .sort_values('AnoMes', ascending=False).head(12).sort_values('AnoMes')
            )
            ultimos_12_meses_df = pd.DataFrame({
                'Mês/Ano': ultimos_12['AnoMes'].dt.strftime('%b/%Y'),
                'Pares Vendidos': ultimos_12['Qtd_Venda'].astype(int),
                'Valor Vendido (R$)': ultimos_12['Vlr_Venda_Corrigido'].map(
                    lambda x: f"R$ {x:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
                )
            })
            st.session_state["ultimos_12_meses_df"] = ultimos_12_meses_df
            st.markdown("### 📆 Vendas dos Últimos 12 Meses")
            st.table(ultimos_12_meses_df)

            # ==== LINHAS/CATEGORIAS NÃO COMPRADAS ====
            st.markdown("## 🧾 Linhas e Categorias que o Cliente Ainda Não Comprou")

            CAMINHO_LINHAS_XLSX = "C:/preditiva/streamlit02/DADOS PREDITIVA LINHAS.xlsx"
            linhas_validas = pd.read_excel(CAMINHO_LINHAS_XLSX, engine="openpyxl")
            linhas_validas.columns = [
                unicodedata.normalize('NFKD', col).encode('ASCII','ignore').decode('utf-8').strip().lower().replace(" ", "_")
                for col in linhas_validas.columns
            ]
            linhas_validas = linhas_validas.drop_duplicates(subset=["linha"])
            linhas_compradas = dados_filtrados.loc[dados_filtrados["Qtd Venda"] > 0, "Linha"].dropna().astype(str)
            linhas_compradas = linhas_compradas.str.strip().str.upper().unique()
            linhas_validas['linha'] = linhas_validas['linha'].astype(str).str.strip().str.upper()
            linhas_validas['codigo_linha'] = linhas_validas['codigo_linha'].astype(str).str.strip()
            linhas_nao_compradas = linhas_validas[~linhas_validas["linha"].isin(linhas_compradas)].copy()
            st.session_state["linhas_nao_compradas"] = linhas_nao_compradas[["codigo_linha","linha"]]

            CAMINHO_CATEGORIAS = "C:/preditiva/streamlit02/CATEGORIAS.csv"
            df_categorias = pd.read_csv(CAMINHO_CATEGORIAS, encoding="latin1", sep=";")
            df_categorias.columns = [
                unicodedata.normalize('NFKD', col).encode('ASCII','ignore').decode('utf-8').strip().lower().replace(" ", "_")
                for col in df_categorias.columns
            ]
            for col in ["categorias","codigo_linha"]:
                if col not in df_categorias.columns:
                    st.error(f"❌ A coluna '{col}' não foi encontrada no arquivo CATEGORIAS.csv.")
                    st.stop()
            df_categorias['categorias'] = df_categorias['categorias'].astype(str).str.strip().str.upper()
            df_categorias['codigo_linha'] = df_categorias['codigo_linha'].astype(str).str.strip()

            linhas_nao_compradas_merge = st.session_state["linhas_nao_compradas"].copy()
            linhas_nao_compradas_merge['codigo_linha'] = linhas_nao_compradas_merge['codigo_linha'].astype(str).str.strip()

            linhas_nao_compradas_categorias = pd.merge(
                linhas_nao_compradas_merge, df_categorias, on="codigo_linha", how="left"
            )
            categorias_nao_compradas = (
                linhas_nao_compradas_categorias[["categorias"]]
                .dropna().drop_duplicates().sort_values(by="categorias").reset_index(drop=True)
            )
            st.session_state["categorias_nao_compradas"] = categorias_nao_compradas

            col1, col2 = st.columns(2)
            with col1:
                st.markdown("### 📄 Linhas que o Cliente Ainda Não Comprou")
                if linhas_nao_compradas.empty:
                    st.success("✅ O cliente comprou todas as linhas.")
                else:
                    st.dataframe(linhas_nao_compradas[["codigo_linha","linha"]].sort_values(by="linha"))
            with col2:
                st.markdown("### 📑 Categorias que o Cliente Ainda Não Comprou")
                if categorias_nao_compradas.empty:
                    st.success("✅ O cliente comprou todas as categorias.")
                else:
                    st.dataframe(categorias_nao_compradas)

            # ==== Persiste no session_state ====
            st.session_state["nome_cliente_para_arquivo"] = nome_grupo.strip().upper().replace(" ","_").replace("/","_")
            st.session_state["nome_grupo"] = nome_grupo
            st.session_state["total_lojas"] = total_lojas
            st.session_state["cods_repr_str"] = cods_repr_str
            st.session_state["nome_supervisor"] = nome_supervisor
            st.session_state["ultima_compra"] = ultima_compra
            st.session_state["periodo_analise"] = periodo_analise
            st.session_state["melhor_mes_nome"] = melhor_mes_nome
            st.session_state["linhas_nao_compradas"] = linhas_nao_compradas
            st.session_state["categorias_nao_compradas"] = categorias_nao_compradas

# ----------------------------
# 📝 EXPORTAÇÃO PARA WORD
# ----------------------------
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import tempfile
from docx import Document

if st.session_state.get("nome_grupo") and st.session_state.get("colecoes_exibir") is not None:
    if st.button("📄 Exportar para Word"):
        with st.spinner("✍️ Gerando documento Word..."):
            nome_grupo = st.session_state["nome_grupo"]
            total_lojas = st.session_state["total_lojas"]
            cods_repr_str = st.session_state["cods_repr_str"]
            nome_supervisor = st.session_state["nome_supervisor"]
            ultima_compra = st.session_state["ultima_compra"]
            periodo_analise = st.session_state["periodo_analise"]
            melhor_mes_nome = st.session_state["melhor_mes_nome"]
            colecoes_exibir = st.session_state["colecoes_exibir"]
            linhas_nao_compradas = st.session_state["linhas_nao_compradas"]
            categorias_nao_compradas = st.session_state["categorias_nao_compradas"]

            def add_heading_colorido(doc, texto, tamanho=14, cor=RGBColor(255, 102, 0)):
                paragrafo = doc.add_paragraph()
                paragrafo.style = None
                run = paragrafo.add_run(texto)
                run.font.bold = True
                run.font.size = Pt(tamanho)
                run.font.color.rgb = cor
                paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            doc = Document()
            paragrafo = doc.add_paragraph()
            paragrafo.style = None
            run = paragrafo.add_run("Relatório Analítico - Kidy")
            run.font.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(255, 102, 0)
            paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            try:
                doc.add_picture("logo_kidy.png", width=Inches(1.5))
            except:
                pass

            doc.add_paragraph(f"📍 Grupo Cliente: {nome_grupo}")
            doc.add_paragraph(f"🏬 Lojas: {total_lojas}")
            doc.add_paragraph(f"🆔 Representante(s): {cods_repr_str}")
            doc.add_paragraph(f"🧑‍💼 Supervisor: {nome_supervisor}")
            doc.add_paragraph(f"📅 Última Compra: {ultima_compra}")
            doc.add_paragraph(f"📊 Período da Análise: {periodo_analise}")
            doc.add_paragraph(f"⭐ Melhor Mês para Oferta: {melhor_mes_nome}")

            add_heading_colorido(doc, "👟 Vendas das 3 Últimas Coleções")
            tabela = doc.add_table(rows=1, cols=4)
            hdr = tabela.rows[0].cells
            hdr[0].text = 'Coleção'
            hdr[1].text = 'Pares Vendidos'
            hdr[2].text = 'Valor Vendido (R$)'
            hdr[3].text = 'Período da Coleta'
            for _, row in colecoes_exibir.iterrows():
                linha = tabela.add_row().cells
                linha[0].text = str(row['Coleção'])
                linha[1].text = str(row['Pares Vendidos'])
                linha[2].text = str(row['Valor Vendido (R$)'])
                linha[3].text = str(row['Período da Coleta'])

            ultimos_12_meses_df = st.session_state.get("ultimos_12_meses_df", pd.DataFrame())
            if not ultimos_12_meses_df.empty and "tabela_12_meses_adicionada" not in st.session_state:
                st.session_state["tabela_12_meses_adicionada"] = True
                add_heading_colorido(doc, "📆 Vendas dos Últimos 12 Meses")
                tabela_12 = doc.add_table(rows=1, cols=3)
                hdr = tabela_12.rows[0].cells
                hdr[0].text = 'Mês/Ano'
                hdr[1].text = 'Pares Vendidos'
                hdr[2].text = 'Valor Vendido (R$)'
                for _, row in ultimos_12_meses_df.iterrows():
                    linha = tabela_12.add_row().cells
                    linha[0].text = str(row['Mês/Ano'])
                    linha[1].text = str(row['Pares Vendidos'])
                    linha[2].text = str(row['Valor Vendido (R$)'])

            add_heading_colorido(doc, "📄 Linhas que o Cliente Ainda Não Comprou")
            if not linhas_nao_compradas.empty:
                for _, row in linhas_nao_compradas.iterrows():
                    doc.add_paragraph(f"- {str(row['codigo_linha'])} | {str(row['linha'])}")
            else:
                doc.add_paragraph("✅ O cliente comprou todas as linhas.")

            add_heading_colorido(doc, "📑 Categorias que o Cliente Ainda Não Comprou")
            if not categorias_nao_compradas.empty:
                for _, row in categorias_nao_compradas.iterrows():
                    doc.add_paragraph(f"- {str(row['categorias'])}")
            else:
                doc.add_paragraph("✅ O cliente comprou todas as categorias.")

            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                doc.save(tmp.name)
                tmp_path = tmp.name

            codigo_repr_para_arquivo = (
                st.session_state.get("codigo_repr_para_arquivo", "representante")
                .replace(" ", "_").lower()
            )
            nome_cliente = st.session_state.get("nome_cliente_para_arquivo", "cliente")
            file_name = f"{nome_cliente}_Rep_{codigo_repr_para_arquivo}.docx"

            with open(tmp_path, "rb") as f:
                st.download_button(
                    label="📥 Baixar Relatório Word",
                    data=f,
                    file_name=file_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

# ----------------------------
# RODAPÉ
# ----------------------------
st.sidebar.markdown("---")
st.sidebar.caption(f"Relatório gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
st.sidebar.markdown("Desenvolvido por Kidy Data Team 🚀")
