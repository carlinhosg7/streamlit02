# ----------------------------
# ✅ IMPORTS
# ----------------------------
import streamlit as st
import pandas as pd
import plotly.express as px
from datetime import datetime
from fpdf import FPDF
import tempfile
import os
from PIL import Image
from sklearn.model_selection import train_test_split
from sklearn.ensemble import RandomForestClassifier
from sklearn.preprocessing import LabelEncoder
from sklearn.metrics import accuracy_score
import time
import plotly.io as pio
import kaleido  # força o loading do kaleido para evitar erro de exportação
import unicodedata
import numpy as np
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import requests
from io import BytesIO
from docx.shared import RGBColor
from docx.shared import Pt
from docx.shared import RGBColor, Pt


pio.kaleido.scope.default_format = "png"

# ----------------------------
# ✅ CONFIGURAÇÃO INICIAL DA PÁGINA (PRIMEIRO COMANDO DO STREAMLIT)
# ----------------------------
st.set_page_config(
    page_title="Dashboard Analítico",
    page_icon="https://raw.githubusercontent.com/carlinhosg7/streamlit02/main/logo_kidy_icon.ico",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ----------------------------
# 🔐 FUNÇÃO DE AUTENTICAÇÃO
# ----------------------------
def autenticar_usuario_excel(caminho_arquivo):
    try:
        df_usuarios = pd.read_excel(caminho_arquivo, engine="openpyxl")

        # Normaliza os nomes das colunas
        df_usuarios.columns = [
            unicodedata.normalize('NFKD', col).encode('ascii', errors='ignore').decode('utf-8').strip().lower().replace(" ", "_")
            for col in df_usuarios.columns
        ]

        # Converte os dados para string
        df_usuarios['usuario'] = df_usuarios['usuario'].apply(lambda x: str(x).strip())
        df_usuarios['senha'] = df_usuarios['senha'].astype(str).str.strip()

        # Inicializa autenticação
        st.session_state['autenticado'] = st.session_state.get('autenticado', False)

        if not st.session_state['autenticado']:
            with st.sidebar:
                st.markdown("### 🔐 Login")
                usuario = st.text_input("Usuário").strip()
                senha = st.text_input("Senha", type="password").strip()

                if st.button("Entrar"):
                    if usuario in df_usuarios['usuario'].values:
                        senha_valida = df_usuarios[df_usuarios['usuario'] == usuario]['senha'].values[0]
                        if senha == senha_valida:
                            st.session_state['autenticado'] = True
                            st.session_state['codigo_representante'] = usuario
                            st.success("✅ Login realizado com sucesso!")
                            st.rerun()
                        else:
                            st.error("❌ Senha incorreta.")
        if not st.session_state['autenticado']:
            st.stop()

    except Exception as e:
        st.error(f"Erro ao carregar planilha de autenticação: {e}")
        st.stop()

# ----------------------------
# 🚀 AUTENTICAÇÃO
# ----------------------------
autenticar_usuario_excel("auth.xlsx")

# ----------------------------
# 🖼️ CARREGA LOGO DA KIDY A PARTIR DO GITHUB
# ----------------------------
url_logo = "https://raw.githubusercontent.com/carlinhosg7/streamlit02/main/logo_kidy.png"
response = requests.get(url_logo)
logo_kidy = Image.open(BytesIO(response.content))

# ----------------------------
# 📊 TÍTULO E SIDEBAR
# ----------------------------
st.title("📊 Dashboard Analítico")
st.sidebar.image(logo_kidy, width=100)
st.sidebar.header("🔧 Filtros de Análise")


# CSS CUSTOMIZADO
def add_custom_css():
    st.markdown("""
        <style>
        body {
            background-color: #1e1e1e;
            color: #ffffff;
        }
        .css-1d391kg {
            background-color: #1e1e1e !important;
        }
        .block-container {
            padding: 2rem;
        }
        div.stButton > button:first-child {
            background-color: #E60012;
            color: white;
            border-radius: 8px;
            height: 3em;
            width: 100%;
            font-weight: bold;
            border: none;
            transition: 0.3s;
        }
        div.stButton > button:first-child:hover {
            background-color: #A3000B;
        }
        footer {visibility: hidden;}

        /* Cards para métricas */
        .metric-card {
            background-color: #2d2d2d;
            padding: 16px;
            border-radius: 10px;
            box-shadow: 1px 1px 8px rgba(0,0,0,0.3);
            text-align: center;
            margin-bottom: 8px;
            border: 1px solid #fba72033;
        }
        .metric-label {
            font-size: 13px;
            color: #bbbbbb;
        }
        .metric-value {
            font-size: 22px;
            font-weight: bold;
            color: #F7A400;
        }
        </style>
    """, unsafe_allow_html=True)

add_custom_css()

# LOGO KIDY (a partir do GitHub)
url_logo = "https://raw.githubusercontent.com/carlinhosg7/streamlit02/main/logo_kidy.png"
response = requests.get(url_logo)
logo_kidy = Image.open(BytesIO(response.content))
st.image(logo_kidy, width=150)


# DICIONÁRIO MESES
meses_portugues = {
    1: 'Janeiro', 2: 'Fevereiro', 3: 'Março', 4: 'Abril',
    5: 'Maio', 6: 'Junho', 7: 'Julho', 8: 'Agosto',
    9: 'Setembro', 10: 'Outubro', 11: 'Novembro', 12: 'Dezembro'
}

# URLS DOS ARQUIVOS
URLS_DADOS = [
    "https://raw.githubusercontent.com/carlinhosg7/streamlit02/main/DADOS_PREDITIVA_1.csv.gz",
    "https://raw.githubusercontent.com/carlinhosg7/streamlit02/main/DADOS_PREDITIVA_2.csv.gz",
    "https://raw.githubusercontent.com/carlinhosg7/streamlit02/main/DADOS_PREDITIVA_3.csv.gz",
    "https://raw.githubusercontent.com/carlinhosg7/streamlit02/main/DADOS_PREDITIVA_4.csv.gz"
]


@st.cache_data(ttl=3600)
def carregar_dados_processados():
    try:
        dfs = [pd.read_csv(url, encoding='cp1252', sep=';', compression='gzip') for url in URLS_DADOS]
        df = pd.concat(dfs, ignore_index=True)
        

        # Converte colunas para tipos eficientes
        df['Codigo Representante'] = df['Codigo Representante'].astype(str).str.strip().str.lstrip("0").astype("category")
        df['Codigo Supervisor'] = df['Codigo Supervisor'].astype(str).str.strip().astype("category")
        df['Codigo Grupo Cliente'] = df['Codigo Grupo Cliente'].astype(str).str.upper().astype("category")
        df['Codigo Cliente'] = df['Codigo Cliente'].astype(str).str.upper().astype("category")
        df['Grupo Cliente'] = df['Grupo Cliente'].astype(str).str.strip().astype("category")
        df['Razao Social'] = df['Razao Social'].astype(str).str.strip().astype("category")
        df['Linha'] = df['Linha'].astype(str).str.strip().astype("category")

        df['Data Cadastro'] = pd.to_datetime(df['Data Cadastro'], errors='coerce')
        df['Data Ultima Compra'] = pd.to_datetime(df['Data Ultima Compra'], errors='coerce')

        df['Qtd Venda'] = pd.to_numeric(df['Qtd Venda'], errors='coerce').fillna(0).astype('int32')
        df['Vlr Venda'] = pd.to_numeric(df['Vlr Venda'], errors='coerce').fillna(0).astype('float32')

        # Cálculo vetorizado sem apply
        df['Preço Médio Produto'] = df['Vlr Venda'] / df['Qtd Venda'].replace(0, pd.NA)
        df['Preço Médio Produto'] = df['Preço Médio Produto'].fillna(0).round(2)
        
            
        # Filtro por representante logado (exceto admin)
        codigo_representante = str(st.session_state.get('codigo_representante', '')).strip().lower().lstrip('0')
        if codigo_representante and codigo_representante != 'admin':
            df = df[df['Codigo Representante'] == codigo_representante]

        return df  # ✅ retorno certo aqui

    except Exception as e:
        import traceback
        st.error("❌ Erro ao carregar dados:")
        st.code(traceback.format_exc())  # Mostra o erro detalhado com linha e tipo
        return pd.DataFrame()

      
        # 🚨 AQUI ENTRA A LÓGICA DE PERFIL ADMIN
        codigo_representante = str(st.session_state.get('codigo_representante', '')).strip().lower().lstrip('0')
        if codigo_representante and codigo_representante != 'admin':
            df = df[df['Codigo Representante'] == codigo_representante]

        return df
    except Exception as e:
        st.error(f"Erro ao carregar dados: {e}")
        return pd.DataFrame()
        # sse bloco vai exibir um aviso claro no topo do dashboard, indicando se o usuário está como admin (vendo tudo) ou mostrando o código do representante comum.
        codigo_representante_logado = st.session_state.get('codigo_representante', '')
        if str(codigo_representante_logado).strip().lower() == 'admin':
            st.markdown("🟢 **Modo Admin: Visualizando todos os dados**")
        else:
            st.markdown(f"🆔 **Representante:** {codigo_representante_logado}")


# CARREGA DADOS
df = carregar_dados_processados()

if df.empty:
    st.stop()

# Opções formatadas para busca com nome
opcoes_grupo_cliente = df[['Codigo Grupo Cliente', 'Grupo Cliente']].drop_duplicates()
opcoes_grupo_cliente['Busca'] = opcoes_grupo_cliente['Codigo Grupo Cliente'].astype(str) + ' - ' + opcoes_grupo_cliente['Grupo Cliente'].astype(str)
busca_grupo = st.sidebar.selectbox("🔍 Buscar Grupo Cliente:", [''] + sorted(opcoes_grupo_cliente['Busca'].tolist()))

opcoes_cliente = df[['Codigo Cliente', 'Razao Social']].drop_duplicates()
opcoes_cliente['Busca'] = opcoes_cliente['Codigo Cliente'].astype(str) + ' - ' + opcoes_cliente['Razao Social'].astype(str)
busca_cliente = st.sidebar.selectbox("🔍 Buscar Cliente:", [''] + sorted(opcoes_cliente['Busca'].tolist()))

# Extrair apenas os códigos selecionados
codigo_grupo_cliente = busca_grupo.split(' - ')[0].strip().upper() if busca_grupo else ''
codigo_cliente = busca_cliente.split(' - ')[0].strip().upper() if busca_cliente else ''

data_min = df['Data Cadastro'].min().date()
data_max = df['Data Cadastro'].max().date()

# 📅 Define período padrão: 01/01/2024 até hoje (limitado ao data_max)
data_inicio_padrao = datetime(2024, 1, 1).date()
data_fim_padrao = datetime.today().date()

# Garante que a data final padrão não ultrapasse data_max
if data_fim_padrao > data_max:
    data_fim_padrao = data_max

periodo = st.sidebar.date_input(
    "Período da análise:",
    value=(data_inicio_padrao, data_fim_padrao),
    min_value=data_min,
    max_value=data_max
)


# BOTÃO
if st.sidebar.button("🔎 Analisar Grupo/Cliente"):

    if not codigo_grupo_cliente and not codigo_cliente:
        st.sidebar.warning("⚠️ Informe pelo menos um código!")
    else:
        with st.spinner('🔎 Analisando dados...'):

            dados_filtrados = df.copy()

            if codigo_cliente:
                dados_filtrados = dados_filtrados[dados_filtrados['Codigo Cliente'] == codigo_cliente]
            elif codigo_grupo_cliente:
                dados_filtrados = dados_filtrados[dados_filtrados['Codigo Grupo Cliente'] == codigo_grupo_cliente]

            dados_filtrados = dados_filtrados[
                (dados_filtrados['Data Cadastro'] >= pd.to_datetime(periodo[0])) &
                (dados_filtrados['Data Cadastro'] <= pd.to_datetime(periodo[1]))
            ]

            if dados_filtrados.empty:
                st.warning("⚠️ Nenhum dado encontrado no período!")
            else:
                dados_filtrados['Ano'] = dados_filtrados['Data Cadastro'].dt.year
                nome_grupo = dados_filtrados['Grupo Cliente'].iloc[0]
                total_lojas = dados_filtrados['Codigo Cliente'].nunique()

                # Mapeamento dos supervisores
                mapa_supervisores = {
                    '9902': 'Centro Oeste',
                    '9907': 'Sul',
                    '9914': 'Norte / Nordeste',
                    '9915': 'REM',
                    '9916': 'SPC',
                    '9917': 'SPI'
                }

                # Extrai representantes únicos do grupo selecionado nos dados filtrados
                cods_representantes = dados_filtrados['Codigo Representante'].astype(str).dropna().unique()
                cods_representantes = sorted(cods_representantes)
                cods_repr_str = ', '.join(cods_representantes)

                # Extrai e limpa código(s) de supervisor
                cod_supervisor = dados_filtrados['Codigo Supervisor'].astype(str).str.strip().dropna().unique()

                # Verifica e exibe o nome do supervisor corretamente
                if len(cod_supervisor) > 0:
                    supervisor_id = str(cod_supervisor[0]).split('.')[0]
                    nome_supervisor = mapa_supervisores.get(supervisor_id, f"Código {supervisor_id} não mapeado")
                else:
                    nome_supervisor = "Não informado"

                # Exibe informações do grupo
                st.markdown(
                    f"## 📍 Grupo Cliente: {nome_grupo} | 🏬 Lojas: {total_lojas} | 🆔 Representante(s): {cods_repr_str}"
                )
                st.markdown(f"#### 🧑‍💼 Supervisor: {nome_supervisor}")

            
                # KPIs
                ultima_data_compra = dados_filtrados['Data Ultima Compra'].max()
                ultima_compra = ultima_data_compra.strftime('%d/%m/%Y') if pd.notnull(ultima_data_compra) else 'Sem compras'

                primeira_data = dados_filtrados['Data Cadastro'].min()
                ultima_data = dados_filtrados['Data Cadastro'].max()
                periodo_analise = f"{primeira_data.strftime('%d/%m/%Y')} até {ultima_data.strftime('%d/%m/%Y')}"

                vendas_totais = dados_filtrados['Qtd Venda'].sum()
                melhor_mes_num = dados_filtrados['Data Cadastro'].dt.month.mode()[0]
                melhor_mes_nome = meses_portugues.get(melhor_mes_num, 'Mês inválido')

                col1, col2, col3 = st.columns(3)
                for col, label, value in zip(
                    [col1, col2, col3],
                    ['📅 Última Compra', '🕒 Período da Análise', '📈 Melhor Mês para Oferta'],
                    [ultima_compra, periodo_analise, melhor_mes_nome]
                ):
                    col.markdown(f"""
                        <div class="metric-card">
                            <div class="metric-label">{label}</div>
                            <div class="metric-value">{value}</div>
                        </div>
                    """, unsafe_allow_html=True)

                st.success(f"📦 Total de Itens Vendidos: {vendas_totais:,} unidades")
    ###################


   ####################             
                # --- ANÁLISE DAS 3 ÚLTIMAS COLEÇÕES (INCLUINDO VIGENTE) ---
                #st.markdown("### 👟 Vendas das 3 Últimas Coleções (Pares e Valores)")

                from datetime import datetime

                # 🔁 Função para identificar a coleção com base na nova regra
                def identificar_colecao(data):
                    if pd.isnull(data):
                        return None
                    mes = data.month
                    ano = data.year

                    if 5 <= mes <= 10:  # Verão: maio a outubro
                        return f"Verão {ano}"
                    elif 11 <= mes <= 12:  # Inverno: novembro e dezembro → pertence ao ano seguinte
                        return f"Inverno {ano + 1}"
                    else:  # Inverno: janeiro a abril → permanece no mesmo ano
                        return f"Inverno {ano}"

                # 📆 Determina a coleção vigente com base na data atual
                hoje = datetime.today()
                mes = hoje.month
                ano = hoje.year

                if 5 <= mes <= 10:
                    colecao_vigente = f"Verão {ano}"
                elif 11 <= mes <= 12:
                    colecao_vigente = f"Inverno {ano + 1}"
                else:
                    colecao_vigente = f"Inverno {ano}"
                
                # 🔧 Garante que os códigos batem corretamente para cliente individual
                df['Codigo Cliente'] = df['Codigo Cliente'].astype(str).str.upper().str.strip()
                codigo_cliente = codigo_cliente.strip().upper()

                # 🔍 Reaplica o filtro por cliente, se estiver definido
                if codigo_cliente:
                    dados_filtrados = dados_filtrados[dados_filtrados['Codigo Cliente'].astype(str).str.upper().str.strip() == codigo_cliente.strip().upper()]





                # Proteção contra dados ausentes
                dados_filtrados = dados_filtrados[
                    pd.notnull(dados_filtrados['Data Cadastro']) &
                    pd.notnull(dados_filtrados['Qtd Venda']) &
                    pd.notnull(dados_filtrados['Preço Médio Produto'])
                ].copy()

                # 👀 Diagnóstico do problema dos valores zerados
                # st.write("🔍 Preço Médio Produto - Quantidade de nulos:", dados_filtrados['Preço Médio Produto'].isnull().sum())
                # st.write("🔍 Preço Médio Produto - Valores únicos:", dados_filtrados['Preço Médio Produto'].unique())


                # Converte corretamente
                # 🔁 Recalcula preço médio e valor corrigido com fallback
                dados_filtrados['Qtd Venda'] = pd.to_numeric(dados_filtrados['Qtd Venda'], errors='coerce').fillna(0)
                dados_filtrados['Vlr Venda'] = pd.to_numeric(dados_filtrados['Vlr Venda'], errors='coerce').fillna(0)

                # Remove registros negativos
                dados_filtrados = dados_filtrados[
                    (dados_filtrados['Qtd Venda'] >= 0) & (dados_filtrados['Vlr Venda'] >= 0)
                ].copy()

                # Preço médio calculado (referência)
                dados_filtrados['Preço Médio Produto'] = dados_filtrados.apply(
                    lambda row: row['Vlr Venda'] / row['Qtd Venda'] if row['Qtd Venda'] > 0 else 0,
                    axis=1
                ).round(2)

                # 🔁 Recalcula campos com proteção total
                dados_filtrados['Qtd Venda'] = pd.to_numeric(dados_filtrados['Qtd Venda'], errors='coerce').fillna(0)
                dados_filtrados['Vlr Venda'] = pd.to_numeric(dados_filtrados['Vlr Venda'], errors='coerce').fillna(0)

                # Remove valores inválidos
                dados_filtrados = dados_filtrados[
                    (dados_filtrados['Qtd Venda'] >= 0) & (dados_filtrados['Vlr Venda'] >= 0)
                ].copy()

                # Preço médio individual por linha (não usado no total, mas útil para verificação)
                dados_filtrados['Preço Médio Produto'] = dados_filtrados.apply(
                    lambda row: row['Vlr Venda'] / row['Qtd Venda'] if row['Qtd Venda'] > 0 else 0,
                    axis=1
                ).round(2)

                # 🔍 Calcula média de preço por par com fallback se necessário
                df_base_preco = dados_filtrados[
                    (dados_filtrados['Qtd Venda'] > 0) & (dados_filtrados['Vlr Venda'] > 0)
                ].copy()

                if not df_base_preco.empty:
                    media_preco_par = df_base_preco.apply(lambda row: row['Vlr Venda'] / row['Qtd Venda'], axis=1).mean()
                else:
                    media_preco_par = 60  # 🔧 Valor estimado padrão por par
                    st.warning(f"⚠️ Nenhuma venda com valor encontrada. Estimativa padrão de R$ {media_preco_par:.2f} por par foi aplicada.")

                # Valor de venda corrigido com fallback
                dados_filtrados['Vlr Venda Corrigido'] = dados_filtrados.apply(
                    lambda row: row['Vlr Venda'] if row['Vlr Venda'] > 0 else row['Qtd Venda'] * media_preco_par,
                    axis=1
                )

                # Define coleção por regra
                def identificar_colecao(data):
                    if pd.isnull(data):
                        return None
                    mes = data.month
                    ano = data.year
                    if 5 <= mes <= 10:
                        return f"Verão {ano}"
                    elif mes >= 11:
                        return f"Inverno {ano + 1}"
                    else:
                        return f"Inverno {ano}"

                dados_filtrados['Colecao'] = dados_filtrados['Data Cadastro'].apply(identificar_colecao)

                # Agrupamento por coleção
                vendas_colecao = dados_filtrados.groupby('Colecao').agg({
                    'Qtd Venda': 'sum',
                    'Vlr Venda Corrigido': 'sum',
                    'Data Cadastro': ['min', 'max']
                }).reset_index()

                vendas_colecao.columns = ['Colecao', 'Qtd Venda', 'Vlr Venda Corrigido', 'Data Inicial', 'Data Final']
                vendas_colecao['Ano'] = vendas_colecao['Colecao'].str.extract(r'(\d{4})').astype(int)

                # Garante coleção vigente
                hoje = datetime.today()
                mes = hoje.month
                ano = hoje.year
                colecao_vigente = f"Verão {ano}" if 5 <= mes <= 10 else f"Inverno {ano + 1}" if mes >= 11 else f"Inverno {ano}"

                if colecao_vigente not in vendas_colecao['Colecao'].values:
                    linha_vigente = pd.DataFrame({
                        'Colecao': [colecao_vigente],
                        'Qtd Venda': [0],
                        'Vlr Venda Corrigido': [0.0],
                        'Data Inicial': [hoje],
                        'Data Final': [hoje],
                        'Ano': [int(colecao_vigente.split()[1])]
                    })
                    vendas_colecao = pd.concat([linha_vigente, vendas_colecao], ignore_index=True)

                # Últimas 3 coleções
                colecoes_exibir = vendas_colecao.sort_values(by='Data Final', ascending=False).drop_duplicates('Colecao').head(3)

                # Formatação final
                colecoes_exibir['Pares Vendidos'] = colecoes_exibir['Qtd Venda'].fillna(0).astype(int)
                colecoes_exibir['Valor Vendido (R$)'] = colecoes_exibir['Vlr Venda Corrigido'].fillna(0).apply(
                    lambda x: f"R$ {x:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
                )
                colecoes_exibir['Período da Coleta'] = colecoes_exibir.apply(
                    lambda row: (
                        f"{row['Data Inicial'].strftime('%d/%m/%Y')} a {row['Data Final'].strftime('%d/%m/%Y')}"
                        if pd.notnull(row['Data Inicial']) and pd.notnull(row['Data Final']) else "Período inválido"
                    ),
                    axis=1
                )

                # Exibe tabela
                colecoes_exibir = colecoes_exibir[['Colecao', 'Pares Vendidos', 'Valor Vendido (R$)', 'Período da Coleta']]
                colecoes_exibir.columns = ['Coleção', 'Pares Vendidos', 'Valor Vendido (R$)', 'Período da Coleta']
                st.markdown("### 👟 Vendas das 3 Últimas Coleções (Pares e Valores)")
                st.table(colecoes_exibir)

                # Salva no session_state para PDF/Word
                st.session_state["colecoes_exibir"] = colecoes_exibir



#########
                # ======================
                # 🧾 Comparativo de Linhas e Categorias - Não Compradas
                # ======================
                import unicodedata  # Para normalização dos nomes das colunas

                st.markdown("## 🧾 Linhas e Categorias que o Cliente Ainda Não Comprou")

               # 🔽 Carrega XLSX com todas as linhas possíveis válidas
                URL_LINHAS_XLSX = "https://github.com/carlinhosg7/streamlit02/raw/main/DADOS%20PREDITIVA%20LINHAS.xlsx"
                linhas_validas = pd.read_excel(URL_LINHAS_XLSX, engine="openpyxl")

                # 🧹 Normaliza nomes das colunas do XLSX
                linhas_validas.columns = [
                    unicodedata.normalize('NFKD', col).encode('ASCII', 'ignore').decode('utf-8').strip().lower().replace(" ", "_")
                    for col in linhas_validas.columns
                ]
                linhas_validas = linhas_validas.drop_duplicates(subset=["linha"])

                # ✅ Linhas compradas pelo cliente
                linhas_compradas = dados_filtrados[dados_filtrados["Qtd Venda"] > 0]["Linha"].dropna().astype(str)
                linhas_compradas = linhas_compradas.str.strip().str.upper().unique()

                # 🔍 Padroniza colunas para comparação
                linhas_validas['linha'] = linhas_validas['linha'].astype(str).str.strip().str.upper()
                linhas_validas['codigo_linha'] = linhas_validas['codigo_linha'].astype(str).str.strip()

                # 🔍 Seleciona as linhas ainda não compradas E válidas
                linhas_nao_compradas = linhas_validas[~linhas_validas["linha"].isin(linhas_compradas)].copy()

                # 🧠 Salva para uso futuro
                st.session_state["linhas_nao_compradas"] = linhas_nao_compradas[["codigo_linha", "linha"]]


                # ================
                # 🔄 Categorias Não Compradas
                # ================
                try:
                    df_categorias = pd.read_csv(
                        "https://raw.githubusercontent.com/carlinhosg7/streamlit02/main/CATEGORIAS.csv",
                        encoding="latin1",
                        sep=";"
                    )
                except Exception as e:
                    st.error(f"Erro ao carregar o arquivo CATEGORIAS.csv: {e}")
                    st.stop()

                # 🧹 Normaliza nomes das colunas
                df_categorias.columns = [
                    unicodedata.normalize('NFKD', col).encode('ASCII', 'ignore').decode('utf-8').strip().lower().replace(" ", "_")
                    for col in df_categorias.columns
                ]

                # ✅ Verifica colunas obrigatórias
                for col in ["categorias", "codigo_linha"]:
                    if col not in df_categorias.columns:
                        st.error(f"❌ A coluna '{col}' não foi encontrada no arquivo CATEGORIAS.csv.")
                        st.stop()

                # 🔧 Padroniza os valores
                df_categorias['categorias'] = df_categorias['categorias'].astype(str).str.strip().str.upper()
                df_categorias['codigo_linha'] = df_categorias['codigo_linha'].astype(str).str.strip()

                # 🔗 Junta com as linhas não compradas
                linhas_nao_compradas_merge = st.session_state["linhas_nao_compradas"].copy()
                linhas_nao_compradas_merge['codigo_linha'] = linhas_nao_compradas_merge['codigo_linha'].astype(str).str.strip()

                linhas_nao_compradas_categorias = pd.merge(
                    linhas_nao_compradas_merge,
                    df_categorias,
                    on="codigo_linha",
                    how="left"
                )

                # 🧽 Categorias únicas
                categorias_nao_compradas = (
                    linhas_nao_compradas_categorias[["categorias"]]
                    .dropna()
                    .drop_duplicates()
                    .sort_values(by="categorias")
                    .reset_index(drop=True)
                )

                # 🧠 Salva para uso no PDF
                st.session_state["categorias_nao_compradas"] = categorias_nao_compradas

                # ======================
                # 📊 Exibição Lado a Lado
                # ======================
                col1, col2 = st.columns(2)

                with col1:
                    st.markdown("### 📄 Linhas que o Cliente Ainda Não Comprou")
                    if linhas_nao_compradas.empty:
                        st.success("✅ O cliente comprou todas as linhas.")
                    else:
                        st.dataframe(linhas_nao_compradas[["codigo_linha", "linha"]].sort_values(by="linha"))

                with col2:
                    st.markdown("### 📑 Categorias que o Cliente Ainda Não Comprou")
                    if categorias_nao_compradas.empty:
                        st.success("✅ O cliente comprou todas as categorias.")
                    else:
                        st.dataframe(categorias_nao_compradas)

                
                
                # 🔐 Salva tudo no session_state no final da análise
                st.session_state["nome_grupo"] = nome_grupo
                st.session_state["total_lojas"] = total_lojas
                st.session_state["cods_repr_str"] = cods_repr_str
                st.session_state["nome_supervisor"] = nome_supervisor
                st.session_state["ultima_compra"] = ultima_compra
                st.session_state["periodo_analise"] = periodo_analise
                st.session_state["melhor_mes_nome"] = melhor_mes_nome
                st.session_state["colecoes_exibir"] = colecoes_exibir
                st.session_state["linhas_nao_compradas"] = linhas_nao_compradas
                st.session_state["categorias_nao_compradas"] = categorias_nao_compradas


                st.session_state["linhas_nao_compradas"] = linhas_nao_compradas
                st.session_state["categorias_nao_compradas"] = categorias_nao_compradas
################



from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import tempfile
from docx import Document

# BOTÃO DE EXPORTAÇÃO PARA WORD
if st.session_state.get("nome_grupo") and st.session_state.get("colecoes_exibir") is not None:

    if st.button("📄 Exportar para Word"):

        with st.spinner("✍️ Gerando documento Word..."):

            # Recupera as variáveis do session_state
            nome_grupo = st.session_state["nome_grupo"]
            total_lojas = st.session_state["total_lojas"]
            cods_repr_str = st.session_state["cods_repr_str"]
            nome_supervisor = st.session_state["nome_supervisor"]
            ultima_compra = st.session_state["ultima_compra"]
            periodo_analise = st.session_state["periodo_analise"]
            melhor_mes_nome = st.session_state["melhor_mes_nome"]
            colecoes_exibir = st.session_state["colecoes_exibir"]
            linhas_nao_compradas = st.session_state["linhas_nao_compradas"]
            categorias_nao_compradas = st.session_state["categorias_nao_compradas"]

            # ✅ Função para adicionar título com cor laranja
            def add_heading_colorido(doc, texto, tamanho=14, cor=RGBColor(255, 102, 0)):
                paragrafo = doc.add_paragraph()
                paragrafo.style = None
                run = paragrafo.add_run(texto)
                run.font.bold = True
                run.font.size = Pt(tamanho)
                run.font.color.rgb = cor
                paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            # Criação do documento Word
            doc = Document()

            paragrafo = doc.add_paragraph()
            paragrafo.style = None  # remove Heading 1 azul do Word
            run = paragrafo.add_run("Relatório Analítico - Kidy")
            run.font.bold = True
            run.font.size = Pt(18)  # Tamanho maior que os subtítulos
            run.font.color.rgb = RGBColor(255, 102, 0)
            paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


            try:
                doc.add_picture("logo_kidy.png", width=Inches(1.5))
            except:
                pass

            doc.add_paragraph(f"📍 Grupo Cliente: {nome_grupo}")
            doc.add_paragraph(f"🏬 Lojas: {total_lojas}")
            doc.add_paragraph(f"🆔 Representante(s): {cods_repr_str}")
            doc.add_paragraph(f"🧑‍💼 Supervisor: {nome_supervisor}")
            doc.add_paragraph(f"📅 Última Compra: {ultima_compra}")
            doc.add_paragraph(f"📊 Período da Análise: {periodo_analise}")
            doc.add_paragraph(f"⭐ Melhor Mês para Oferta: {melhor_mes_nome}")

            # Tabela de Coleções
            add_heading_colorido(doc, "👟 Vendas das 3 Últimas Coleções")
            tabela = doc.add_table(rows=1, cols=4)
            hdr = tabela.rows[0].cells
            hdr[0].text = 'Coleção'
            hdr[1].text = 'Pares Vendidos'
            hdr[2].text = 'Valor Vendido (R$)'
            hdr[3].text = 'Período da Coleta'

            for _, row in colecoes_exibir.iterrows():
                linha = tabela.add_row().cells
                linha[0].text = str(row['Coleção'])
                linha[1].text = str(row['Pares Vendidos'])
                linha[2].text = str(row['Valor Vendido (R$)'])
                linha[3].text = str(row['Período da Coleta'])

            # Linhas não compradas
            add_heading_colorido(doc, "📄 Linhas que o Cliente Ainda Não Comprou")
            if not linhas_nao_compradas.empty:
                for _, row in linhas_nao_compradas.iterrows():
                    doc.add_paragraph(f"- {row['linha']}")
            else:
                doc.add_paragraph("✅ O cliente comprou todas as linhas.")

            # Categorias não compradas
            add_heading_colorido(doc, "📑 Categorias que o Cliente Ainda Não Comprou")
            if not categorias_nao_compradas.empty:
                for _, row in categorias_nao_compradas.iterrows():
                    doc.add_paragraph(f"- {row['categorias']}")
            else:
                doc.add_paragraph("✅ O cliente comprou todas as categorias.")

            # Exportação final
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                doc.save(tmp.name)
                tmp_path = tmp.name

            with open(tmp_path, "rb") as f:
                st.download_button(
                    label="📥 Baixar Relatório Word",
                    data=f,
                    file_name="relatorio_analitico_kidy.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

# RODAPÉ
st.sidebar.markdown("---")
st.sidebar.caption(f"Relatório gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
st.sidebar.markdown("Desenvolvido por Kidy Data Team 🚀")
