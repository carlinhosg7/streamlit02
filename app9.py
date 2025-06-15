import streamlit as st
import pandas as pd
import plotly.express as px
from datetime import datetime
from fpdf import FPDF
import tempfile
import os
from PIL import Image
from sklearn.model_selection import train_test_split
from sklearn.ensemble import RandomForestClassifier
from sklearn.preprocessing import LabelEncoder
from sklearn.metrics import accuracy_score
import time
import plotly.io as pio
import kaleido # força o loading do kaleido para evitar erro de importação
pio.kaleido.scope.default_format = "png"
import plotly.express as px
import unicodedata
import numpy as np
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import requests
from io import BytesIO


# 🔐 FUNÇÃO DE AUTENTICAÇÃO
def autenticar_usuario_excel(caminho_arquivo):
    try:
        df_usuarios = pd.read_excel(caminho_arquivo, engine="openpyxl")

        # Normaliza os nomes das colunas
        df_usuarios.columns = [
            unicodedata.normalize('NFKD', col).encode('ascii', errors='ignore').decode('utf-8').strip().lower().replace(" ", "_")
            for col in df_usuarios.columns
        ]

        # Converte os dados para string
        df_usuarios['usuario'] = df_usuarios['usuario'].apply(lambda x: str(x).strip())
        df_usuarios['senha'] = df_usuarios['senha'].astype(str).str.strip()

        # Inicializa autenticação
        st.session_state['autenticado'] = st.session_state.get('autenticado', False)

        if not st.session_state['autenticado']:
            with st.sidebar:
                st.markdown("### 🔐 Login")
                usuario = st.text_input("Usuário").strip()
                senha = st.text_input("Senha", type="password").strip()

                if st.button("Entrar"):
                    if usuario in df_usuarios['usuario'].values:
                        senha_valida = df_usuarios[df_usuarios['usuario'] == usuario]['senha'].values[0]
                        if senha == senha_valida:
                            st.session_state['autenticado'] = True
                            st.session_state['codigo_representante'] = usuario  # <-- Adicione esta linha
                            st.success("✅ Login realizado com sucesso!")
                            st.rerun()
                        else:
                            st.error("❌ Senha incorreta.")  

        if not st.session_state['autenticado']:
            st.stop()

    except Exception as e:
        st.error(f"Erro ao carregar planilha de autenticação: {e}")
        st.stop()

# 🎨 CONFIG PÁGINA (PRIMEIRA CHAMADA DO STREAMLIT)
st.set_page_config(
    page_title="Análise de Clientes",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 🚀 SÓ DEPOIS AUTENTICAÇÃO

# 🔗 Caminho do arquivo auth.xlsx no GitHub
URL_AUTH = "https://raw.githubusercontent.com/carlinhosg7/streamlit02/main/auth.xlsx"

# ⬇️ Baixar e carregar o arquivo em memória
try:
    response = requests.get(URL_AUTH)
    response.raise_for_status()  # dispara erro se não encontrar

    auth_file = BytesIO(response.content)

    # 🧪 Chamar função com arquivo em memória
    autenticar_usuario_excel(auth_file)

except Exception as e:
    st.error(f"Erro ao baixar o arquivo de autenticação: {e}")
    st.stop()



st.title("🔓 Login feito com sucesso, Análise Liberada")

# CSS CUSTOMIZADO
def add_custom_css():
    st.markdown("""
        <style>
        body {
            background-color: #1e1e1e;
            color: #ffffff;
        }
        .css-1d391kg {
            background-color: #1e1e1e !important;
        }
        .block-container {
            padding: 2rem;
        }
        div.stButton > button:first-child {
            background-color: #E60012;
            color: white;
            border-radius: 8px;
            height: 3em;
            width: 100%;
            font-weight: bold;
            border: none;
            transition: 0.3s;
        }
        div.stButton > button:first-child:hover {
            background-color: #A3000B;
        }
        footer {visibility: hidden;}

        /* Cards para métricas */
        .metric-card {
            background-color: #2d2d2d;
            padding: 16px;
            border-radius: 10px;
            box-shadow: 1px 1px 8px rgba(0,0,0,0.3);
            text-align: center;
            margin-bottom: 8px;
            border: 1px solid #fba72033;
        }
        .metric-label {
            font-size: 13px;
            color: #bbbbbb;
        }
        .metric-value {
            font-size: 22px;
            font-weight: bold;
            color: #F7A400;
        }
        </style>
    """, unsafe_allow_html=True)

add_custom_css()

# LOGO KIDY
logo_kidy = Image.open("logo_kidy.png")
st.image(logo_kidy, width=150)

# DICIONÁRIO MESES
meses_portugues = {
    1: 'Janeiro', 2: 'Fevereiro', 3: 'Março', 4: 'Abril',
    5: 'Maio', 6: 'Junho', 7: 'Julho', 8: 'Agosto',
    9: 'Setembro', 10: 'Outubro', 11: 'Novembro', 12: 'Dezembro'
}

# URLS DOS ARQUIVOS
URLS_DADOS = [
    "https://raw.githubusercontent.com/carlinhosg7/streamlit02/main/DADOS_PREDITIVA_1.csv.gz",
    "https://raw.githubusercontent.com/carlinhosg7/streamlit02/main/DADOS_PREDITIVA_2.csv.gz",
    "https://raw.githubusercontent.com/carlinhosg7/streamlit02/main/DADOS_PREDITIVA_3.csv.gz",
    "https://raw.githubusercontent.com/carlinhosg7/streamlit02/main/DADOS_PREDITIVA_4.csv.gz"
]


@st.cache_data(ttl=3600)
def carregar_dados_processados():
    try:
        dfs = [pd.read_csv(url, encoding='cp1252', sep=';', compression='gzip') for url in URLS_DADOS]
        df = pd.concat(dfs, ignore_index=True)
        

        # Converte colunas para tipos eficientes
        df['Codigo Representante'] = df['Codigo Representante'].astype(str).str.strip().str.lstrip("0").astype("category")
        df['Codigo Supervisor'] = df['Codigo Supervisor'].astype(str).str.strip().astype("category")
        df['Codigo Grupo Cliente'] = df['Codigo Grupo Cliente'].astype(str).str.upper().astype("category")
        df['Codigo Cliente'] = df['Codigo Cliente'].astype(str).str.upper().astype("category")
        df['Grupo Cliente'] = df['Grupo Cliente'].astype(str).str.strip().astype("category")
        df['Razao Social'] = df['Razao Social'].astype(str).str.strip().astype("category")
        df['Linha'] = df['Linha'].astype(str).str.strip().astype("category")

        df['Data Cadastro'] = pd.to_datetime(df['Data Cadastro'], errors='coerce')
        df['Data Ultima Compra'] = pd.to_datetime(df['Data Ultima Compra'], errors='coerce')

        df['Qtd Venda'] = pd.to_numeric(df['Qtd Venda'], errors='coerce').fillna(0).astype('int32')
        df['Vlr Venda'] = pd.to_numeric(df['Vlr Venda'], errors='coerce').fillna(0).astype('float32')

        # Cálculo vetorizado sem apply
        df['Preço Médio Produto'] = df['Vlr Venda'] / df['Qtd Venda'].replace(0, pd.NA)
        df['Preço Médio Produto'] = df['Preço Médio Produto'].fillna(0).round(2)
        
            
        # Filtro por representante logado (exceto admin)
        codigo_representante = str(st.session_state.get('codigo_representante', '')).strip().lower().lstrip('0')
        if codigo_representante and codigo_representante != 'admin':
            df = df[df['Codigo Representante'] == codigo_representante]

        return df  # ✅ retorno certo aqui

    except Exception as e:
        import traceback
        st.error("❌ Erro ao carregar dados:")
        st.code(traceback.format_exc())  # Mostra o erro detalhado com linha e tipo
        return pd.DataFrame()

      
        # 🚨 AQUI ENTRA A LÓGICA DE PERFIL ADMIN
        codigo_representante = str(st.session_state.get('codigo_representante', '')).strip().lower().lstrip('0')
        if codigo_representante and codigo_representante != 'admin':
            df = df[df['Codigo Representante'] == codigo_representante]

        return df
    except Exception as e:
        st.error(f"Erro ao carregar dados: {e}")
        return pd.DataFrame()
        # sse bloco vai exibir um aviso claro no topo do dashboard, indicando se o usuário está como admin (vendo tudo) ou mostrando o código do representante comum.
        codigo_representante_logado = st.session_state.get('codigo_representante', '')
        if str(codigo_representante_logado).strip().lower() == 'admin':
            st.markdown("🟢 **Modo Admin: Visualizando todos os dados**")
        else:
            st.markdown(f"🆔 **Representante:** {codigo_representante_logado}")


# CARREGA DADOS
df = carregar_dados_processados()

if df.empty:
    st.stop()

# FILTROS SIDEBAR
st.title("📊 Dashboard Analítico")
st.sidebar.image(logo_kidy, width=100)
st.sidebar.header("🔧 Filtros de Análise")

# Opções formatadas para busca com nome
opcoes_grupo_cliente = df[['Codigo Grupo Cliente', 'Grupo Cliente']].drop_duplicates()
opcoes_grupo_cliente['Busca'] = opcoes_grupo_cliente['Codigo Grupo Cliente'].astype(str) + ' - ' + opcoes_grupo_cliente['Grupo Cliente'].astype(str)
busca_grupo = st.sidebar.selectbox("🔍 Buscar Grupo Cliente:", [''] + sorted(opcoes_grupo_cliente['Busca'].tolist()))

opcoes_cliente = df[['Codigo Cliente', 'Razao Social']].drop_duplicates()
opcoes_cliente['Busca'] = opcoes_cliente['Codigo Cliente'].astype(str) + ' - ' + opcoes_cliente['Razao Social'].astype(str)
busca_cliente = st.sidebar.selectbox("🔍 Buscar Cliente:", [''] + sorted(opcoes_cliente['Busca'].tolist()))

# Extrair apenas os códigos selecionados
codigo_grupo_cliente = busca_grupo.split(' - ')[0].strip().upper() if busca_grupo else ''
codigo_cliente = busca_cliente.split(' - ')[0].strip().upper() if busca_cliente else ''

data_min = df['Data Cadastro'].min().date()
data_max = df['Data Cadastro'].max().date()

# 📅 Define período padrão: 01/01/2024 até hoje (limitado ao data_max)
data_inicio_padrao = datetime(2024, 1, 1).date()
data_fim_padrao = datetime.today().date()

# Garante que a data final padrão não ultrapasse data_max
if data_fim_padrao > data_max:
    data_fim_padrao = data_max

periodo = st.sidebar.date_input(
    "Período da análise:",
    value=(data_inicio_padrao, data_fim_padrao),
    min_value=data_min,
    max_value=data_max
)


# BOTÃO
if st.sidebar.button("🔎 Analisar Grupo/Cliente"):

    if not codigo_grupo_cliente and not codigo_cliente:
        st.sidebar.warning("⚠️ Informe pelo menos um código!")
    else:
        with st.spinner('🔎 Analisando dados...'):

            dados_filtrados = df.copy()

            if codigo_cliente:
                dados_filtrados = dados_filtrados[dados_filtrados['Codigo Cliente'] == codigo_cliente]
            elif codigo_grupo_cliente:
                dados_filtrados = dados_filtrados[dados_filtrados['Codigo Grupo Cliente'] == codigo_grupo_cliente]

            dados_filtrados = dados_filtrados[
                (dados_filtrados['Data Cadastro'] >= pd.to_datetime(periodo[0])) &
                (dados_filtrados['Data Cadastro'] <= pd.to_datetime(periodo[1]))
            ]

            if dados_filtrados.empty:
                st.warning("⚠️ Nenhum dado encontrado no período!")
            else:
                dados_filtrados['Ano'] = dados_filtrados['Data Cadastro'].dt.year
                nome_grupo = dados_filtrados['Grupo Cliente'].iloc[0]
                total_lojas = dados_filtrados['Codigo Cliente'].nunique()

                # Mapeamento dos supervisores
                mapa_supervisores = {
                    '9902': 'Centro Oeste',
                    '9907': 'Sul',
                    '9914': 'Norte / Nordeste',
                    '9915': 'REM',
                    '9916': 'SPC',
                    '9917': 'SPI'
                }

                # Extrai representantes únicos do grupo selecionado nos dados filtrados
                cods_representantes = dados_filtrados['Codigo Representante'].astype(str).dropna().unique()
                cods_representantes = sorted(cods_representantes)
                cods_repr_str = ', '.join(cods_representantes)

                # Extrai e limpa código(s) de supervisor
                cod_supervisor = dados_filtrados['Codigo Supervisor'].astype(str).str.strip().dropna().unique()

                # Verifica e exibe o nome do supervisor corretamente
                if len(cod_supervisor) > 0:
                    supervisor_id = str(cod_supervisor[0]).split('.')[0]
                    nome_supervisor = mapa_supervisores.get(supervisor_id, f"Código {supervisor_id} não mapeado")
                else:
                    nome_supervisor = "Não informado"

                # Exibe informações do grupo
                st.markdown(
                    f"## 📍 Grupo Cliente: {nome_grupo} | 🏬 Lojas: {total_lojas} | 🆔 Representante(s): {cods_repr_str}"
                )
                st.markdown(f"#### 🧑‍💼 Supervisor: {nome_supervisor}")

            
                # KPIs
                ultima_data_compra = dados_filtrados['Data Ultima Compra'].max()
                ultima_compra = ultima_data_compra.strftime('%d/%m/%Y') if pd.notnull(ultima_data_compra) else 'Sem compras'

                primeira_data = dados_filtrados['Data Cadastro'].min()
                ultima_data = dados_filtrados['Data Cadastro'].max()
                periodo_analise = f"{primeira_data.strftime('%d/%m/%Y')} até {ultima_data.strftime('%d/%m/%Y')}"

                vendas_totais = dados_filtrados['Qtd Venda'].sum()
                melhor_mes_num = dados_filtrados['Data Cadastro'].dt.month.mode()[0]
                melhor_mes_nome = meses_portugues.get(melhor_mes_num, 'Mês inválido')

                col1, col2, col3 = st.columns(3)
                for col, label, value in zip(
                    [col1, col2, col3],
                    ['📅 Última Compra', '🕒 Período da Análise', '📈 Melhor Mês para Oferta'],
                    [ultima_compra, periodo_analise, melhor_mes_nome]
                ):
                    col.markdown(f"""
                        <div class="metric-card">
                            <div class="metric-label">{label}</div>
                            <div class="metric-value">{value}</div>
                        </div>
                    """, unsafe_allow_html=True)

                st.success(f"📦 Total de Itens Vendidos: {vendas_totais:,} unidades")
    ###################


   ####################             
                # --- ANÁLISE DAS 3 ÚLTIMAS COLEÇÕES (INCLUINDO VIGENTE) ---
                #st.markdown("### 👟 Vendas das 3 Últimas Coleções (Pares e Valores)")

                from datetime import datetime

                # 🔁 Função para identificar a coleção com base na nova regra
                def identificar_colecao(data):
                    if pd.isnull(data):
                        return None
                    mes = data.month
                    ano = data.year

                    if 5 <= mes <= 10:  # Verão: maio a outubro
                        return f"Verão {ano}"
                    elif 11 <= mes <= 12:  # Inverno: novembro e dezembro → pertence ao ano seguinte
                        return f"Inverno {ano + 1}"
                    else:  # Inverno: janeiro a abril → permanece no mesmo ano
                        return f"Inverno {ano}"

                # 📆 Determina a coleção vigente com base na data atual
                hoje = datetime.today()
                mes = hoje.month
                ano = hoje.year

                if 5 <= mes <= 10:
                    colecao_vigente = f"Verão {ano}"
                elif 11 <= mes <= 12:
                    colecao_vigente = f"Inverno {ano + 1}"
                else:
                    colecao_vigente = f"Inverno {ano}"
                
                # 🔧 Garante que os códigos batem corretamente para cliente individual
                df['Codigo Cliente'] = df['Codigo Cliente'].astype(str).str.upper().str.strip()
                codigo_cliente = codigo_cliente.strip().upper()

                # 🔍 Reaplica o filtro por cliente, se estiver definido
                if codigo_cliente:
                    dados_filtrados = dados_filtrados[dados_filtrados['Codigo Cliente'].astype(str).str.upper().str.strip() == codigo_cliente.strip().upper()]





                # Proteção contra dados ausentes
                dados_filtrados = dados_filtrados[
                    pd.notnull(dados_filtrados['Data Cadastro']) &
                    pd.notnull(dados_filtrados['Qtd Venda']) &
                    pd.notnull(dados_filtrados['Preço Médio Produto'])
                ].copy()

                # 👀 Diagnóstico do problema dos valores zerados
                # st.write("🔍 Preço Médio Produto - Quantidade de nulos:", dados_filtrados['Preço Médio Produto'].isnull().sum())
                # st.write("🔍 Preço Médio Produto - Valores únicos:", dados_filtrados['Preço Médio Produto'].unique())


                # Converte para numérico
                dados_filtrados['Qtd Venda'] = pd.to_numeric(dados_filtrados['Qtd Venda'], errors='coerce').fillna(0).astype(float)
                dados_filtrados['Preço Médio Produto'] = pd.to_numeric(dados_filtrados['Preço Médio Produto'], errors='coerce').fillna(0).astype(float)

                # Valor corrigido
                dados_filtrados['Vlr Venda Corrigido'] = dados_filtrados['Qtd Venda'] * dados_filtrados['Preço Médio Produto']

                # Define coleção
                def identificar_colecao(data):
                    if pd.isnull(data):
                        return None
                    mes = data.month
                    ano = data.year
                    if 5 <= mes <= 10:
                        return f"Verão {ano}"
                    elif mes >= 11:
                        return f"Inverno {ano + 1}"
                    else:
                        return f"Inverno {ano}"

                dados_filtrados['Colecao'] = dados_filtrados['Data Cadastro'].apply(identificar_colecao)

                # Agrupamento
                vendas_colecao = dados_filtrados.groupby('Colecao').agg({
                    'Qtd Venda': 'sum',
                    'Vlr Venda Corrigido': 'sum',
                    'Data Cadastro': ['min', 'max']
                }).reset_index()

                # Renomeia colunas
                vendas_colecao.columns = ['Colecao', 'Qtd Venda', 'Vlr Venda Corrigido', 'Data Inicial', 'Data Final']
                vendas_colecao['Data Inicial'] = pd.to_datetime(vendas_colecao['Data Inicial'], errors='coerce')
                vendas_colecao['Data Final'] = pd.to_datetime(vendas_colecao['Data Final'], errors='coerce')
                vendas_colecao['Ano'] = vendas_colecao['Colecao'].str.extract(r'(\d{4})').astype(int)

                # Garante coleção vigente
                if colecao_vigente not in vendas_colecao['Colecao'].values:
                    linha_vigente = pd.DataFrame({
                        'Colecao': [colecao_vigente],
                        'Qtd Venda': [0],
                        'Vlr Venda Corrigido': [0.0],
                        'Data Inicial': [hoje],
                        'Data Final': [hoje],
                        'Ano': [int(colecao_vigente.split()[1])]
                    })
                    vendas_colecao = pd.concat([linha_vigente, vendas_colecao], ignore_index=True)

                # 3 últimas coleções
                colecoes_exibir = vendas_colecao.sort_values(by='Data Final', ascending=False).drop_duplicates('Colecao').head(3)

                # Formatações
                colecoes_exibir['Pares Vendidos'] = colecoes_exibir['Qtd Venda'].fillna(0).astype(int)
                colecoes_exibir['Valor Vendido (R$)'] = colecoes_exibir['Vlr Venda Corrigido'].fillna(0).apply(
                    lambda x: f"R$ {x:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
                )
                colecoes_exibir['Período da Coleta'] = colecoes_exibir.apply(
                    lambda row: (
                        f"{row['Data Inicial'].strftime('%d/%m/%Y')} a {row['Data Final'].strftime('%d/%m/%Y')}"
                        if pd.notnull(row['Data Inicial']) and pd.notnull(row['Data Final'])
                        else "Período inválido"
                    ),
                    axis=1
                )

                # Exibição
                colecoes_exibir = colecoes_exibir[['Colecao', 'Pares Vendidos', 'Valor Vendido (R$)', 'Período da Coleta']]
                colecoes_exibir.columns = ['Coleção', 'Pares Vendidos', 'Valor Vendido (R$)', 'Período da Coleta']
                st.markdown("### 👟 Vendas das 3 Últimas Coleções (Pares e Valores)")
                st.table(colecoes_exibir)

#########
                # ======================
                # 🧾 Comparativo de Linhas e Categorias - Não Compradas
                # ======================
                import unicodedata  # Para normalização dos nomes das colunas

                st.markdown("## 🧾 Linhas e Categorias que o Cliente Ainda Não Comprou")

               # 🔽 Carrega XLSX com todas as linhas possíveis válidas
                URL_LINHAS_XLSX = "https://github.com/carlinhosg7/streamlit02/raw/main/DADOS%20PREDITIVA%20LINHAS.xlsx"
                linhas_validas = pd.read_excel(URL_LINHAS_XLSX, engine="openpyxl")

                # 🧹 Normaliza nomes das colunas do XLSX
                linhas_validas.columns = [
                    unicodedata.normalize('NFKD', col).encode('ASCII', 'ignore').decode('utf-8').strip().lower().replace(" ", "_")
                    for col in linhas_validas.columns
                ]
                linhas_validas = linhas_validas.drop_duplicates(subset=["linha"])

                # ✅ Linhas compradas pelo cliente
                linhas_compradas = dados_filtrados[dados_filtrados["Qtd Venda"] > 0]["Linha"].dropna().astype(str)
                linhas_compradas = linhas_compradas.str.strip().str.upper().unique()

                # 🔍 Padroniza colunas para comparação
                linhas_validas['linha'] = linhas_validas['linha'].astype(str).str.strip().str.upper()
                linhas_validas['codigo_linha'] = linhas_validas['codigo_linha'].astype(str).str.strip()

                # 🔍 Seleciona as linhas ainda não compradas E válidas
                linhas_nao_compradas = linhas_validas[~linhas_validas["linha"].isin(linhas_compradas)].copy()

                # 🧠 Salva para uso futuro
                st.session_state["linhas_nao_compradas"] = linhas_nao_compradas[["codigo_linha", "linha"]]


                # ================
                # 🔄 Categorias Não Compradas
                # ================
                try:
                    df_categorias = pd.read_csv(
                        "https://raw.githubusercontent.com/carlinhosg7/streamlit02/main/CATEGORIAS.csv",
                        encoding="latin1",
                        sep=";"
                    )
                except Exception as e:
                    st.error(f"Erro ao carregar o arquivo CATEGORIAS.csv: {e}")
                    st.stop()

                # 🧹 Normaliza nomes das colunas
                df_categorias.columns = [
                    unicodedata.normalize('NFKD', col).encode('ASCII', 'ignore').decode('utf-8').strip().lower().replace(" ", "_")
                    for col in df_categorias.columns
                ]

                # ✅ Verifica colunas obrigatórias
                for col in ["categorias", "codigo_linha"]:
                    if col not in df_categorias.columns:
                        st.error(f"❌ A coluna '{col}' não foi encontrada no arquivo CATEGORIAS.csv.")
                        st.stop()

                # 🔧 Padroniza os valores
                df_categorias['categorias'] = df_categorias['categorias'].astype(str).str.strip().str.upper()
                df_categorias['codigo_linha'] = df_categorias['codigo_linha'].astype(str).str.strip()

                # 🔗 Junta com as linhas não compradas
                linhas_nao_compradas_merge = st.session_state["linhas_nao_compradas"].copy()
                linhas_nao_compradas_merge['codigo_linha'] = linhas_nao_compradas_merge['codigo_linha'].astype(str).str.strip()

                linhas_nao_compradas_categorias = pd.merge(
                    linhas_nao_compradas_merge,
                    df_categorias,
                    on="codigo_linha",
                    how="left"
                )

                # 🧽 Categorias únicas
                categorias_nao_compradas = (
                    linhas_nao_compradas_categorias[["categorias"]]
                    .dropna()
                    .drop_duplicates()
                    .sort_values(by="categorias")
                    .reset_index(drop=True)
                )

                # 🧠 Salva para uso no PDF
                st.session_state["categorias_nao_compradas"] = categorias_nao_compradas

                # ======================
                # 📊 Exibição Lado a Lado
                # ======================
                col1, col2 = st.columns(2)

                with col1:
                    st.markdown("### 📄 Linhas que o Cliente Ainda Não Comprou")
                    if linhas_nao_compradas.empty:
                        st.success("✅ O cliente comprou todas as linhas.")
                    else:
                        st.dataframe(linhas_nao_compradas[["codigo_linha", "linha"]].sort_values(by="linha"))

                with col2:
                    st.markdown("### 📑 Categorias que o Cliente Ainda Não Comprou")
                    if categorias_nao_compradas.empty:
                        st.success("✅ O cliente comprou todas as categorias.")
                    else:
                        st.dataframe(categorias_nao_compradas)

                
                
                # 🔐 Salva tudo no session_state no final da análise
                st.session_state["nome_grupo"] = nome_grupo
                st.session_state["total_lojas"] = total_lojas
                st.session_state["cods_repr_str"] = cods_repr_str
                st.session_state["nome_supervisor"] = nome_supervisor
                st.session_state["ultima_compra"] = ultima_compra
                st.session_state["periodo_analise"] = periodo_analise
                st.session_state["melhor_mes_nome"] = melhor_mes_nome
                st.session_state["colecoes_exibir"] = colecoes_exibir
                st.session_state["linhas_nao_compradas"] = linhas_nao_compradas
                st.session_state["categorias_nao_compradas"] = categorias_nao_compradas


                st.session_state["linhas_nao_compradas"] = linhas_nao_compradas
                st.session_state["categorias_nao_compradas"] = categorias_nao_compradas
################

                # MACHINE LEARNING
                st.subheader("🤖 Previsão de Linhas para Oferta (Machine Learning) Exeto Botas e Coturnos")

                # Inicializa a barra de progresso
                progress_bar = st.progress(0, text="⏳ Preparando dados para ML (0%)")

                # Etapa 1: preparar dados
                @st.cache_data(ttl=600)
                def preparar_dados_ml(df):
                    df_ml = df.copy()
                    df_ml['Data Cadastro'] = pd.to_datetime(df_ml['Data Cadastro'], errors='coerce')
                    df_ml['Mes Pedido'] = df_ml['Data Cadastro'].dt.month
                    df_ml['Compra'] = df_ml['Qtd Venda'].apply(lambda x: 1 if x > 0 else 0)
                    return df_ml

                dados_ml = preparar_dados_ml(df)
                progress_bar.progress(20, text="✅ Etapa 1: Dados preparados (20%)")

                # Etapa 2: Treinar modelo
                @st.cache_resource
                def treinar_modelo_rf(df_ml):
                    from sklearn.ensemble import RandomForestClassifier
                    from sklearn.model_selection import train_test_split
                    from sklearn.metrics import accuracy_score
                    from sklearn.preprocessing import LabelEncoder
                    import numpy as np

                    df_ml = df_ml.copy()

                    le_grupo = LabelEncoder().fit(df_ml['Codigo Grupo Cliente'])
                    le_cliente = LabelEncoder().fit(df_ml['Codigo Cliente'])
                    le_linha = LabelEncoder().fit(df_ml['Linha'])

                    df_ml['Grupo_Code'] = le_grupo.transform(df_ml['Codigo Grupo Cliente'])
                    df_ml['Cliente_Code'] = le_cliente.transform(df_ml['Codigo Cliente'])
                    df_ml['Linha_Code'] = le_linha.transform(df_ml['Linha'])

                    X = df_ml[['Grupo_Code', 'Cliente_Code', 'Linha_Code', 'Mes Pedido']]
                    y = df_ml['Compra']

                    X = X.replace([np.inf, -np.inf], np.nan)
                    y = y.replace([np.inf, -np.inf], np.nan)
                    dados_validos = X.notna().all(axis=1) & y.notna()
                    X = X[dados_validos]
                    y = y[dados_validos]

                    X_train, X_test, y_train, y_test = train_test_split(
                        X, y, test_size=0.2, stratify=y, random_state=42
                    )

                    modelo = RandomForestClassifier(n_estimators=100, random_state=42)
                    modelo.fit(X_train, y_train)
                    acc = accuracy_score(y_test, modelo.predict(X_test))

                    return modelo, le_grupo, le_cliente, le_linha, acc

                modelo_rf, le_grupo, le_cliente, le_linha, acc = treinar_modelo_rf(dados_ml)
                progress_bar.progress(50, text="✅ Etapa 2: Modelo treinado (50%)")

                # Etapa 3: Predição
                grupo_id = codigo_grupo_cliente or dados_filtrados['Codigo Grupo Cliente'].iloc[0]
                cliente_id = codigo_cliente or dados_filtrados['Codigo Cliente'].iloc[0]
                mes_atual = datetime.now().month

                # 1️⃣ Carrega planilha de categorias
                URL_CATEGORIAS = "https://raw.githubusercontent.com/carlinhosg7/streamlit02/main/CATEGORIAS.csv"
                df_categorias = pd.read_csv(URL_CATEGORIAS, sep=";", encoding="latin1")

                # 2️⃣ Normaliza colunas
                df_categorias.columns = (
                    df_categorias.columns
                    .str.strip()
                    .str.lower()
                    .str.replace(" ", "_")
                    .str.normalize('NFKD')
                    .str.encode('ascii', errors='ignore')
                    .str.decode('utf-8')
                )
                # Remove decimais e converte para string limpa
                df_categorias["codigo_linha"] = df_categorias["codigo_linha"].apply(lambda x: str(x).split('.')[0].strip() if pd.notnull(x) else "")

                # 3️⃣ Prepara linhas possíveis
                linhas_possiveis = df[['Linha', 'Codigo Linha']].dropna().drop_duplicates()
                linhas_possiveis['Codigo Linha'] = linhas_possiveis['Codigo Linha'].apply(lambda x: str(int(x)).strip() if pd.notnull(x) else "")

                # 4️⃣ Merge com categorias (apenas uma vez!)
                linhas_possiveis = linhas_possiveis.merge(
                    df_categorias,
                    left_on='Codigo Linha',
                    right_on='codigo_linha',
                    how='left'
                )

                # 5️⃣ Remove categorias proibidas
                categorias_bloqueadas = ['BOTA', 'COTURNO']
                coluna_categoria = next((col for col in linhas_possiveis.columns if 'categoria' in col.lower()), None)

                if coluna_categoria:
                    linhas_possiveis = linhas_possiveis[~linhas_possiveis[coluna_categoria].str.upper().isin(categorias_bloqueadas)]
                else:
                    st.warning("⚠️ Coluna 'categoria' não encontrada após o merge. Verifique o nome no CSV.")

                try:
                    # 6️⃣ Prepara dados para predição
                    dados_para_prever = pd.DataFrame({
                        'Grupo_Code': le_grupo.transform([grupo_id] * len(linhas_possiveis)),
                        'Cliente_Code': le_cliente.transform([cliente_id] * len(linhas_possiveis)),
                        'Linha_Code': le_linha.transform(linhas_possiveis['Linha']),
                        'Mes Pedido': [mes_atual] * len(linhas_possiveis)
                    })

                    progress_bar.progress(70, text="✅ Etapa 3: Dados para predição gerados (70%)")

                    # 7️⃣ Predição
                    probs = modelo_rf.predict_proba(dados_para_prever)[:, 1].ravel()

                    df_preds = pd.DataFrame({
                        'Código da Linha': linhas_possiveis['Codigo Linha'].values,
                        'Linha': linhas_possiveis['Linha'].values,
                        'Probabilidade de Compra': probs,
                        'Categoria': linhas_possiveis[coluna_categoria].values if coluna_categoria else [''] * len(linhas_possiveis)
                    }).sort_values(by='Probabilidade de Compra', ascending=False)

                    # Limpa formatação do código da linha
                    df_preds['Código da Linha'] = df_preds['Código da Linha'].apply(lambda x: str(x).split('.')[0])

                    # 💾 Salva no session_state
                    st.session_state["df_preds"] = df_preds

                    progress_bar.progress(100, text="✅ Etapa 4: Predição finalizada (100%)")
                    st.success("🎉 Predição realizada com sucesso!")

                    df_preds_formatado = df_preds.head(10).copy()
                    df_preds_formatado['Probabilidade de Compra'] = df_preds_formatado['Probabilidade de Compra'].apply(
                        lambda x: f"{x:.0%}".replace('.', ',')
                    )

                    st.table(df_preds_formatado[['Código da Linha', 'Linha', 'Categoria', 'Probabilidade de Compra']])

                except ValueError as e:
                    st.warning(f"⚠️ Erro ao prever: {e}. Verifique se o código do cliente ou grupo existe nos dados.")
                    progress_bar.progress(0, text="❌ Erro durante a predição.")







                st.session_state["df_preds"] = df_preds

                # Garante que o vetor de probabilidades está 1D
                probs = probs.ravel()

                # Monta o DataFrame de previsão
                ddf_preds = pd.DataFrame({
                    'Código da Linha': linhas_possiveis['Codigo Linha'].values,
                    'Linha': linhas_possiveis['Linha'].values,
                    'Probabilidade de Compra': probs,
                    'Categoria': linhas_possiveis[coluna_categoria].values if coluna_categoria else [''] * len(linhas_possiveis)
                }).sort_values(by='Probabilidade de Compra', ascending=False)


                # Remove os números depois do ponto no Código da Linha
                df_preds['Código da Linha'] = df_preds['Código da Linha'].apply(
                    lambda x: str(int(float(x))) if pd.notnull(x) else ""
                )






                # Salvar no session_state para exportação
                st.session_state["df_preds"] = df_preds



# BOTÃO DE EXPORTAÇÃO PARA WORD

if st.session_state.get("nome_grupo") and st.session_state.get("colecoes_exibir") is not None:

    if st.button("📄 Exportar para Word"):

        with st.spinner("✍️ Gerando documento Word..."):

            # Recupera as variáveis do session_state
            nome_grupo = st.session_state["nome_grupo"]
            total_lojas = st.session_state["total_lojas"]
            cods_repr_str = st.session_state["cods_repr_str"]
            nome_supervisor = st.session_state["nome_supervisor"]
            ultima_compra = st.session_state["ultima_compra"]
            periodo_analise = st.session_state["periodo_analise"]
            melhor_mes_nome = st.session_state["melhor_mes_nome"]
            colecoes_exibir = st.session_state["colecoes_exibir"]
            linhas_nao_compradas = st.session_state["linhas_nao_compradas"]
            categorias_nao_compradas = st.session_state["categorias_nao_compradas"]

            # Criação do documento Word
            doc = Document()
            doc.add_heading("Relatório Analítico - Kidy", level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            try:
                doc.add_picture("logo_kidy.png", width=Inches(1.5))
            except:
                pass

            doc.add_paragraph(f"📍 Grupo Cliente: {nome_grupo}")
            doc.add_paragraph(f"🏬 Lojas: {total_lojas}")
            doc.add_paragraph(f"🆔 Representante(s): {cods_repr_str}")
            doc.add_paragraph(f"🧑‍💼 Supervisor: {nome_supervisor}")
            doc.add_paragraph(f"📅 Última Compra: {ultima_compra}")
            doc.add_paragraph(f"📊 Período da Análise: {periodo_analise}")
            doc.add_paragraph(f"⭐ Melhor Mês para Oferta: {melhor_mes_nome}")

            # Tabela de Coleções
            doc.add_heading("👟 Vendas das 3 Últimas Coleções", level=2)
            tabela = doc.add_table(rows=1, cols=4)
            tabela.style = 'Table Grid'  # Garante bordas visíveis

            # Cabeçalho
            hdr = tabela.rows[0].cells
            hdr[0].text = 'Coleção'
            hdr[1].text = 'Pares Vendidos'
            hdr[2].text = 'Valor Vendido (R$)'
            hdr[3].text = 'Período da Coleta'

            # Linhas da tabela
            for _, row in colecoes_exibir.iterrows():
                linha = tabela.add_row().cells
                linha[0].text = str(row['Coleção'])

                # Formata pares vendidos com separador de milhar
                linha[1].text = f"{int(row['Pares Vendidos']):,}".replace(",", ".")

                # Formata valor vendido no estilo brasileiro (ex: R$ 203.406,00)
                valor_bruto = str(row['Valor Vendido (R$)']).replace("R$", "").replace(".", "").replace(",", ".").strip()

                try:
                    valor_float = float(valor_bruto)
                    linha[2].text = f"R$ {valor_float:,.2f}".replace(",", "#").replace(".", ",").replace("#", ".")
                except ValueError:
                    linha[2].text = "R$ 0,00"


                # Período da coleta sem alteração
                linha[3].text = str(row['Período da Coleta'])



            # Linhas não compradas
            doc.add_heading("📄 Linhas que o Cliente Ainda Não Comprou", level=2)
            if not linhas_nao_compradas.empty:
                for _, row in linhas_nao_compradas.iterrows():
                    doc.add_paragraph(f"- {row['linha']}")
            else:
                doc.add_paragraph("✅ O cliente comprou todas as linhas.")

            # Categorias não compradas
            doc.add_heading("📑 Categorias que o Cliente Ainda Não Comprou", level=2)
            if not categorias_nao_compradas.empty:
                for _, row in categorias_nao_compradas.iterrows():
                    doc.add_paragraph(f"- {row['categorias']}")
            else:
                doc.add_paragraph("✅ O cliente comprou todas as categorias.")
            
           
            # 🔍 Tabela de Previsão ML (se existir)
            if "df_preds" in st.session_state and not st.session_state["df_preds"].empty:
                df_preds = st.session_state["df_preds"].head(10)

                doc.add_heading("🧠 Previsão de Linhas para Oferta (Machine Learning) Exeto Botas e Coturnos)", level=2)

                # Cria a tabela com 3 colunas fixas
                tabela_ml = doc.add_table(rows=1, cols=3)
                tabela_ml.style = 'Table Grid'

                # Cabeçalhos
                hdr_cells = tabela_ml.rows[0].cells
                hdr_cells[0].text = "Código da Linha"
                hdr_cells[1].text = "Linha"
                hdr_cells[2].text = "Probabilidade de Compra"

                # Dados formatados
                for _, row in df_preds.iterrows():
                    row_cells = tabela_ml.add_row().cells
                    row_cells[0].text = str(row["Código da Linha"])
                    row_cells[1].text = str(row["Linha"])
                    row_cells[2].text = f"{row['Probabilidade de Compra']:.0%}".replace('.', ',')
            else:
                doc.add_paragraph("⚠️ Previsão de compra indisponível ou sem dados suficientes.")



            # Exportação final
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
                doc.save(tmp.name)
                tmp_path = tmp.name

            with open(tmp_path, "rb") as f:
                st.download_button(
                    label="📥 Baixar Relatório Word",
                    data=f,
                    file_name="relatorio_analitico_kidy.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

# RODAPÉ
st.sidebar.markdown("---")
st.sidebar.caption(f"Relatório gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
st.sidebar.markdown("Desenvolvido por Kidy Data Team 🚀")
